from docx import Document

from utils import _CAPTION_NUM_PATTERN
import json

from validate_docx import (
    check_assets_completeness,
    check_center_misjudgment,
    check_figure_table_integrity,
    check_long_bold_run,
    count_docx_formulas,
)


def test_caption_pattern_exposes_number_group():
    match = _CAPTION_NUM_PATTERN.match("表 12 描述性统计")
    assert match and match.group(1) == "12"


def test_figure_table_integrity_does_not_crash_on_caption(tmp_path):
    doc = Document()
    doc.add_paragraph("表1 描述性统计")
    doc.add_paragraph("图1 分布图")
    path = tmp_path / "paper.docx"
    doc.save(path)
    result = check_figure_table_integrity(str(path), None, None, None)
    assert result["table_nums"] == [1]
    assert result["figure_nums"] == [1]


def test_minimal_docx_has_no_formulas(tmp_path):
    doc = Document()
    doc.add_paragraph("正文")
    path = tmp_path / "paper.docx"
    doc.save(path)
    assert count_docx_formulas(str(path)) == (0, 0)


def test_assets_resolve_relative_to_manifest(tmp_path):
    output = tmp_path / "output"
    (output / "tables").mkdir(parents=True)
    (output / "tables/t.md").write_text("|x|\n|-|\n|1|", encoding="utf-8")
    manifest = output / "assets_manifest.json"
    manifest.write_text(json.dumps({
        "tables": [{"id": "t", "path": "tables/t.md", "required": True}],
        "figures": [],
    }), encoding="utf-8")
    doc = Document()
    doc.add_table(rows=1, cols=1)
    paper = tmp_path / "paper.docx"
    doc.save(paper)
    result = check_assets_completeness(str(paper), None, str(manifest))
    assert result["status"] == "pass"


def test_first_paragraph_is_allowed_as_centered_bold_title(tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("这是一个长度超过三十个汉字并且应当正常居中加粗显示的完整论文标题示例文本")
    run.bold = True
    doc.add_paragraph("正文")
    path = tmp_path / "paper.docx"
    doc.save(path)
    assert check_long_bold_run(str(path))["status"] == "pass"
    assert check_center_misjudgment(str(path))["status"] == "pass"
