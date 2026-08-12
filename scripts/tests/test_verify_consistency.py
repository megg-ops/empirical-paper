import json

from docx import Document

from verify_consistency import _check_citations, _check_structured_numbers, verify_paper


def _results(path):
    path.write_text(json.dumps({"reportable_values": [{
        "key": "model.beta", "value_raw": 1.2344, "value_display": "1.234",
        "allowed_text_forms": ["1.234", "1.23"], "must_report": True,
    }]}), encoding="utf-8")


def test_must_report_accepts_declared_precision_only():
    assert _check_structured_numbers("系数为 1.23。", {
        "reportable_values": [{"key": "b", "value_raw": 1.234, "allowed_text_forms": ["1.23"], "must_report": True}]
    })["status"] == "PASS"
    assert _check_structured_numbers("系数为 1.24。", {
        "reportable_values": [{"key": "b", "value_raw": 1.234, "allowed_text_forms": ["1.23"], "must_report": True}]
    })["status"] == "BLOCKER"


def test_citation_gap_is_blocker():
    result = _check_citations("正文[1][3]\n# 参考文献\n[1] A\n[3] C")
    assert result["status"] == "BLOCKER"


def test_docx_table_cells_participate_in_number_check(tmp_path):
    paper = tmp_path / "paper.docx"
    doc = Document()
    doc.add_paragraph("正文")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "1.234"
    doc.save(paper)
    results = tmp_path / "results.json"
    _results(results)
    result = verify_paper(str(paper), str(results), paper_format="docx", skip_word_count=True)
    assert result["checks"]["structured_numbers"]["status"] == "PASS"


def test_missing_input_returns_structured_blocker(tmp_path):
    result = verify_paper(str(tmp_path / "none.md"), str(tmp_path / "none.json"), skip_word_count=True)
    assert result["verdict"] == "FAIL"
