#!/usr/bin/env python3
"""
validate_docx.py -- Word 文件验证脚本

检查 docx 文件的可打开性、公式完整性、图表完整性、表格结构、字体段落、引用参考文献。

用法：
    python validate_docx.py \
        --docx paper_workspace/final_paper/paper_final.docx \
        --markdown paper_workspace/04_writer/output/paper_draft.md \
        --tables paper_workspace/03_coder/output/tables \
        --figures paper_workspace/03_coder/output/figures \
        --output paper_workspace/final_paper/docx_validation_report.md

Exit codes:
    0: no BLOCKER (may have WARN)
    2: has BLOCKER
"""

import argparse
import logging
import os
import re
import sys
import zipfile
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from lxml import etree
except ImportError:
    etree = None

from utils import (
    count_md_formulas,
    extract_md_citation_numbers,
    _CAPTION_NUM_PATTERN,
    _CAPTION_REF_VERB_PATTERN,
)
from quality_contract import make_quality_result, write_json


# ==================== Check A: Openability ====================

def check_openability(docx_path: str) -> dict:
    """检查 docx 文件可打开性"""
    result = {'status': 'pass', 'issues': []}

    # 1. zipfile
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            names = zf.namelist()
            if 'word/document.xml' not in names:
                result['status'] = 'block'
                result['issues'].append('word/document.xml 不存在于 zip 中')
                return result
            doc_xml_bytes = zf.read('word/document.xml')
    except zipfile.BadZipFile:
        result['status'] = 'block'
        result['issues'].append('文件不是有效的 zip/docx')
        return result
    except Exception as e:
        result['status'] = 'block'
        result['issues'].append(f'zip 打开失败: {e}')
        return result

    # 2. python-docx
    try:
        from docx import Document
        Document(docx_path)
    except ImportError:
        result['issues'].append('python-docx 未安装，跳过 docx 读取检查')
        result['status'] = 'warn'
    except Exception as e:
        result['status'] = 'block'
        result['issues'].append(f'python-docx 无法读取: {e}')
        return result

    # 3. XML 解析
    if etree is not None:
        try:
            etree.fromstring(doc_xml_bytes)
        except etree.XMLSyntaxError as e:
            result['status'] = 'block'
            result['issues'].append(f'document.xml 解析失败: {e}')

    return result


# ==================== Check B: Formula Integrity ====================

def count_docx_formulas(docx_path: str) -> tuple[int, int]:
    """统计 docx 中的 oMath 和 oMathPara 数量"""
    omath_count = 0
    omathpara_count = 0

    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            try:
                doc_xml = zf.read('word/document.xml').decode('utf-8')
            except Exception:
                return 0, 0
    except Exception:
        return 0, 0

    if etree is None:
        # 纯文本匹配
        omath_count = len(re.findall(r'<m:oMath[ >]', doc_xml))
        omathpara_count = len(re.findall(r'<m:oMathPara[ >]', doc_xml))
    else:
        try:
            root = etree.fromstring(doc_xml.encode('utf-8') if isinstance(doc_xml, str) else doc_xml)
            ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
            omath_count = len(root.findall('.//m:oMath', ns))
            omathpara_count = len(root.findall('.//m:oMathPara', ns))
        except Exception:
            omath_count = len(re.findall(r'<m:oMath[ >]', doc_xml))
            omathpara_count = len(re.findall(r'<m:oMathPara[ >]', doc_xml))

    return omath_count, omathpara_count


def check_formula_integrity(docx_path: str, md_path: str) -> dict:
    """检查公式完整性"""
    result = {
        'status': 'pass', 'issues': [],
        'md_inline': 0, 'md_block': 0,
        'docx_omath': 0, 'docx_omathpara': 0,
        'latex_residue': [], 'loss_sentences': [],
    }

    # Markdown 公式统计
    md_text = ''
    if md_path and os.path.exists(md_path):
        md_text = Path(md_path).read_text(encoding='utf-8')
        inline, block = count_md_formulas(md_text)
        result['md_inline'] = inline
        result['md_block'] = block

    # DOCX 公式统计
    omath, omathpara = count_docx_formulas(docx_path)
    result['docx_omath'] = omath
    result['docx_omathpara'] = omathpara

    # 提取 docx 文本用于后续检查
    doc_text = ''
    try:
        from docx import Document
        doc = Document(docx_path)
        doc_text = '\n'.join([p.text for p in doc.paragraphs])
    except Exception:
        pass

    # LaTeX 残留检查
    latex_patterns = [
        r'\\alpha\b', r'\\beta\b', r'\\varepsilon\b', r'\\sum\b',
        r'\\frac\b', r'\\theta\b', r'\\lambda\b', r'\\sqrt\b',
        r'\\int\b', r'\\prod\b', r'\\left\b', r'\\right\b',
    ]
    for pat in latex_patterns:
        matches = re.findall(pat, doc_text)
        if matches:
            result['latex_residue'].extend([f'{pat}: {len(matches)} 处'])

    if len(result['latex_residue']) > 3:
        result['status'] = 'block'
        result['issues'].append(f'大量 LaTeX 命令残留: {len(result["latex_residue"])} 种')

    # 公式变量缺失句式
    loss_patterns = [
        r'设共有个', r'其中，为', r'若，说明',
        r'设共有\s*个决策单元', r'其中[，,]\s*为',
    ]
    for pat in loss_patterns:
        matches = re.findall(pat, doc_text)
        if matches:
            result['loss_sentences'].extend(matches)

    if result['loss_sentences']:
        result['status'] = 'block'
        result['issues'].append(f'公式变量缺失句式: {result["loss_sentences"]}')

    # Markdown 有独立公式但 docx 没有任何公式对象
    if result['md_block'] > 0 and omath == 0 and omathpara == 0:
        result['status'] = 'block'
        result['issues'].append(
            f'Markdown 有 {result["md_block"]} 个独立公式，但 docx 中无任何 oMath/oMathPara'
        )

    # Word 公式数少于 Markdown（WARN）
    if result['md_block'] > 0 and (omath + omathpara) > 0:
        md_total = result['md_inline'] + result['md_block']
        docx_total = omath + omathpara
        if docx_total < md_total * 0.8 and result['status'] != 'block':
            result['status'] = 'warn'
            result['issues'].append(
                f'Word 公式数 ({docx_total}) 少于 Markdown 公式数 ({md_total})'
            )

    return result


# ==================== Check C: Figure/Table Integrity ====================

def check_figure_table_integrity(docx_path: str, md_path: str,
                                  tables_dir: str = None,
                                  figures_dir: str = None) -> dict:
    """检查图表完整性"""
    result = {'status': 'pass', 'issues': [],
              'md_images': 0, 'docx_images': 0,
              'table_nums': [], 'figure_nums': []}

    doc_text = ''
    try:
        from docx import Document
        doc = Document(docx_path)
        doc_text = '\n'.join([p.text for p in doc.paragraphs])
    except Exception:
        pass

    # Markdown 图片路径检查
    if md_path and os.path.exists(md_path):
        md_text = Path(md_path).read_text(encoding='utf-8')
        img_paths = re.findall(r'!\[.*?\]\(([^)]+)\)', md_text)
        result['md_images'] = len(img_paths)

        md_dir = os.path.dirname(md_path)
        for img in img_paths:
            full = os.path.join(md_dir, img) if not os.path.isabs(img) else img
            if not os.path.exists(full):
                result['issues'].append(f'Markdown 图片路径不存在: {img}')
                result['status'] = 'block'

    # DOCX 嵌入图片数
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            result['docx_images'] = len([n for n in zf.namelist()
                                         if n.startswith('word/media/')])
    except Exception:
        pass

    # Markdown vs DOCX 图片数对比
    if result['md_images'] > 0 and result['docx_images'] < result['md_images']:
        result['status'] = 'block'
        result['issues'].append(
            f'DOCX 嵌入图片数不足：Markdown {result["md_images"]} 张，DOCX {result["docx_images"]} 张'
        )

    # 图表编号连续性
    # 只检查真正的标题段落（"表X 标题内容"格式），排除正文引用（"表X报告了..."）
    caption_num_pattern = _CAPTION_NUM_PATTERN
    ref_verb_pattern = _CAPTION_REF_VERB_PATTERN
    table_nums = []
    figure_nums = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m = caption_num_pattern.match(text)
        if not m:
            continue
        if ref_verb_pattern.match(text):
            continue  # 正文引用，不是标题
        if text.startswith('表'):
            table_nums.append(int(m.group(1)))
        elif text.startswith('图'):
            figure_nums.append(int(m.group(1)))
    table_nums = sorted(set(table_nums))
    figure_nums = sorted(set(figure_nums))
    result['table_nums'] = table_nums
    result['figure_nums'] = figure_nums

    for label, nums in [('表', table_nums), ('图', figure_nums)]:
        if nums:
            expected = list(range(1, max(nums) + 1))
            missing = sorted(set(expected) - set(nums))
            duplicates = [n for n in nums if nums.count(n) > 1]
            if missing:
                result['issues'].append(f'{label}编号不连续: 缺{missing}')
                result['status'] = 'block'
            if duplicates:
                result['issues'].append(f'{label}编号重复: {sorted(set(duplicates))}')
                result['status'] = 'block'
            if nums != sorted(nums):
                result['issues'].append(f'{label}编号乱序: {nums}')
                result['status'] = 'block'

    return result


# ==================== Check D: Table Structure ====================

def check_table_structure(docx_path: str) -> dict:
    """检查表格结构"""
    result = {'status': 'pass', 'issues': [], 'table_count': 0, 'table_details': []}

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception:
        result['issues'].append('无法读取 docx，跳过表格检查')
        result['status'] = 'warn'
        return result

    tables = doc.tables
    result['table_count'] = len(tables)

    for i, table in enumerate(tables):
        nrows = len(table.rows)
        ncols = len(table.columns) if table.rows else 0
        result['table_details'].append(f'表{i+1}: {nrows}行×{ncols}列')

        # 1列超长表格
        if ncols == 1 and nrows > 10:
            result['status'] = 'block'
            result['issues'].append(f'表{i+1} 仅1列但{nrows}行，可能被压成纯文本')

        # 全空表格
        all_empty = all(
            cell.text.strip() == ''
            for row in table.rows
            for cell in row.cells
        )
        if all_empty:
            result['status'] = 'block'
            result['issues'].append(f'表{i+1} 所有单元格为空')

        # 宽表风险检查
        ncols = len(table.columns)
        if ncols >= 10:
            result['status'] = 'block'
            result['issues'].append(f'表{i+1} 有 {ncols} 列（≥10），未拆表/转图/横向页，可能导致排版溢出')
        elif ncols >= 8:
            if result['status'] == 'pass':
                result['status'] = 'warn'
            result['issues'].append(f'表{i+1} 有 {ncols} 列（≥8），较宽，建议检查排版')

        # 重复 tcBorders 检查
        try:
            from docx.oxml.ns import qn
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        borders_list = tcPr.findall(qn('w:tcBorders'))
                        if len(borders_list) > 1:
                            result['status'] = 'block'
                            result['issues'].append(
                                f'表{i+1} 存在重复 tcBorders（{len(borders_list)}个）'
                            )
                            break
                else:
                    continue
                break
        except Exception:
            pass

    return result


# ==================== Check E: Font & Paragraph ====================

def check_font_paragraph(docx_path: str) -> dict:
    """检查字体和段落（仅 WARN）"""
    result = {'status': 'pass', 'issues': []}

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(docx_path)
    except Exception:
        result['issues'].append('无法读取 docx，跳过字体检查')
        return result

    # 检查 Heading 样式
    heading_count = 0
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith('Heading'):
            heading_count += 1

    if heading_count == 0:
        result['issues'].append('未找到 Heading 样式，可能标题层级未正确设置')
        result['status'] = 'warn'

    # 检查中文字体 eastAsia 设置（抽样）
    eastasia_missing = 0
    for p in doc.paragraphs[:50]:
        for run in p.runs:
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'))
                    if not ea:
                        eastasia_missing += 1

    if eastasia_missing > 10:
        result['issues'].append(f'有 {eastasia_missing} 个 run 缺少 eastAsia 字体设置')
        result['status'] = 'warn'

    return result


# ==================== Check F: Citations & References ====================

def check_citations_references(docx_path: str) -> dict:
    """检查引用和参考文献"""
    result = {'status': 'pass', 'issues': [],
              'body_citations': [], 'reference_citations': []}

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception:
        result['issues'].append('无法读取 docx，跳过引用检查')
        return result

    full_text = '\n'.join([p.text for p in doc.paragraphs])

    # 提取引用编号
    body_nums = set(extract_md_citation_numbers(full_text))

    result['body_citations'] = sorted(body_nums)

    # 查找参考文献部分
    ref_nums = set()
    in_refs = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if '参考文献' in text and len(text) < 20:
            in_refs = True
            continue
        if in_refs:
            for m in re.finditer(r'\[(\d+)\]', text):
                ref_nums.add(int(m.group(1)))

    result['reference_citations'] = sorted(ref_nums)

    # 正文引用无对应参考文献
    if body_nums and ref_nums:
        missing_refs = sorted(body_nums - ref_nums)
        if missing_refs:
            result['status'] = 'block'
            result['issues'].append(f'正文引用 [{missing_refs}] 在参考文献中不存在')

        # 参考文献未被正文引用
        unused_refs = sorted(ref_nums - body_nums)
        if unused_refs:
            result['issues'].append(f'参考文献 [{unused_refs}] 未被正文引用')
            if result['status'] == 'pass':
                result['status'] = 'warn'

    return result


# ==================== Check G: Long Bold Run ====================

def check_long_bold_run(docx_path: str) -> dict:
    """检查正文中异常长的加粗段落"""
    result = {'status': 'pass', 'issues': [], 'details': []}

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception:
        result['issues'].append('无法读取 docx，跳过长加粗检查')
        return result

    # 收集表格内段落（用于排除）
    table_paras = set()
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        table_paras.add(id(p))
    except Exception:
        pass

    first_content = next((p for p in doc.paragraphs if p.text.strip()), None)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 文档首个非空段落是论文标题；居中加粗属于标题正常格式。
        if first_content is not None and para._p is first_content._p:
            continue
        # 跳过标题
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            continue
        # 跳过表格内段落
        if id(para) in table_paras:
            continue
        # 跳过摘要/关键词标签
        if text.startswith('摘要') or text.startswith('关键词'):
            continue

        # 统计连续加粗的汉字数
        bold_chars = 0
        for run in para.runs:
            if run.bold:
                bold_chars += len(re.findall(r'[\u4e00-\u9fff]', run.text))
            else:
                bold_chars = 0  # 连续中断

        # 也用整体方式检查：如果所有 bold run 合并后汉字数
        total_bold = 0
        in_bold = False
        max_run = 0
        current_run = 0
        for run in para.runs:
            if run.bold:
                cn = len(re.findall(r'[\u4e00-\u9fff]', run.text))
                if in_bold:
                    current_run += cn
                else:
                    in_bold = True
                    current_run = cn
            else:
                if in_bold:
                    max_run = max(max_run, current_run)
                in_bold = False
                current_run = 0
        if in_bold:
            max_run = max(max_run, current_run)

        if max_run >= 80:
            result['status'] = 'block'
            preview = text[:60] + '...' if len(text) > 60 else text
            result['issues'].append(f'超长加粗段落（{max_run}汉字）: {preview}')
        elif max_run >= 30:
            preview = text[:60] + '...' if len(text) > 60 else text
            result['issues'].append(f'较长加粗段落（{max_run}汉字）: {preview}')
            if result['status'] == 'pass':
                result['status'] = 'warn'

    return result


# ==================== Check H: Center Misjudgment ====================

def check_center_misjudgment(docx_path: str) -> dict:
    """检查 Normal 样式段落被误居中（特别是正文引用被当成图表标题）"""
    result = {'status': 'pass', 'issues': [], 'details': []}

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document(docx_path)
    except Exception:
        result['issues'].append('无法读取 docx，跳过居中误判检查')
        return result

    # 收集表格内段落（用于排除）
    table_paras = set()
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        table_paras.add(id(p))
    except Exception:
        pass

    # 引用模式：表1报告了 / 图1展示了...
    ref_pattern = _CAPTION_REF_VERB_PATTERN
    # 真标题模式
    caption_pattern = _CAPTION_NUM_PATTERN

    first_content = next((p for p in doc.paragraphs if p.text.strip()), None)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if first_content is not None and para._p is first_content._p:
            continue
        # 只检查居中的段落
        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            continue
        # 跳过标题
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            continue
        # 跳过非 Normal 样式
        if para.style and para.style.name and para.style.name != 'Normal':
            continue
        # 跳过表格内段落
        if id(para) in table_paras:
            continue
        # 跳过摘要/关键词标签
        if text.startswith('摘要') or text.startswith('关键词'):
            continue
        # 跳过真正的图表标题
        if caption_pattern.match(text) and not ref_pattern.match(text):
            continue

        # 到这里：Normal + 居中 + 不是真正标题 → 可疑
        preview = text[:80] + '...' if len(text) > 80 else text
        if ref_pattern.match(text):
            # 正文引用被居中 → BLOCKER
            result['status'] = 'block'
            result['issues'].append(f'正文引用被误居中: {preview}')
        else:
            # 其他 Normal 居中 → WARN
            result['issues'].append(f'Normal 样式段落被居中: {preview}')
            if result['status'] == 'pass':
                result['status'] = 'warn'

    return result


# ==================== Check J: Assets Completeness ====================

def check_assets_completeness(docx_path: str, md_path: str, assets_manifest_path: str) -> dict:
    """检查资产占位符是否全部替换，required 资产是否嵌入"""
    result = {'status': 'pass', 'issues': [], 'stats': {}}

    if not assets_manifest_path or not os.path.exists(assets_manifest_path):
        result['status'] = 'pass'
        result['stats'] = {'skipped': True, 'reason': 'no assets_manifest provided'}
        return result

    import json
    try:
        with open(assets_manifest_path, 'r', encoding='utf-8') as f:
            assets = json.load(f)
    except Exception as e:
        result['status'] = 'warn'
        result['issues'].append(f'assets_manifest 读取失败: {e}')
        return result

    manifest_dir = Path(assets_manifest_path).resolve().parent

    def asset_path(item: dict) -> Path:
        path = Path(item['path'])
        return path if path.is_absolute() else manifest_dir / path

    fig_map = {f['id']: f for f in assets.get('figures', [])}
    tbl_map = {t['id']: t for t in assets.get('tables', [])}

    # Read docx text for placeholder check
    try:
        from docx import Document
        doc = Document(docx_path)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
    except Exception:
        full_text = ''

    # Check for leftover [FIGURE: xxx] placeholders
    leftover_fig = re.findall(r'\[FIGURE:\s*(\w+)\]', full_text)
    leftover_tbl = re.findall(r'\[TABLE:\s*(\w+)\]', full_text)

    if leftover_fig:
        result['status'] = 'block'
        result['issues'].append(f'文中有未替换的 [FIGURE:] 占位符: {leftover_fig}')
    if leftover_tbl:
        result['status'] = 'block'
        result['issues'].append(f'文中有未替换的 [TABLE:] 占位符: {leftover_tbl}')

    # Check for leftover TABLE_ASSET comments in docx XML
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            doc_xml = zf.read('word/document.xml').decode('utf-8', errors='replace')
        leftover_asset_comments = re.findall(r'TABLE_ASSET:([^:]+):', doc_xml)
        if leftover_asset_comments:
            result['status'] = 'warn'
            result['issues'].append(f'docx XML 中有未处理的 TABLE_ASSET 注释: {leftover_asset_comments}')
    except Exception:
        pass

    # Count required assets
    required_figs = [f for f in assets.get('figures', []) if f.get('required', True)]
    required_tbls = [t for t in assets.get('tables', []) if t.get('required', True)]

    # Check that required figure files exist
    missing_fig_files = []
    for fig in required_figs:
        if not asset_path(fig).exists():
            missing_fig_files.append(fig['id'])
    if missing_fig_files:
        result['status'] = 'block'
        result['issues'].append(f'required 图片文件不存在: {missing_fig_files}')

    # Check that required table files exist
    missing_tbl_files = []
    for tbl in required_tbls:
        if not asset_path(tbl).exists():
            missing_tbl_files.append(tbl['id'])
    if missing_tbl_files:
        result['status'] = 'block'
        result['issues'].append(f'required 表格文件不存在: {missing_tbl_files}')

    # Count images in docx
    docx_images = 0
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            docx_images = len([n for n in zf.namelist() if n.startswith('word/media/')])
    except Exception:
        pass

    # Check actual table count in docx vs required tables
    docx_tables = 0
    try:
        from docx import Document as _Doc
        _d = _Doc(docx_path)
        docx_tables = len(_d.tables)
    except Exception:
        pass

    required_tables_count = len(required_tbls)
    if required_tables_count > 0 and docx_tables < required_tables_count:
        result['status'] = 'block'
        result['issues'].append(
            f'required_tables={required_tables_count} 但 docx 实际表格数={docx_tables}，表格缺失'
        )

    result['stats'] = {
        'assets_manifest': assets_manifest_path,
        'required_figures': len(required_figs),
        'required_tables': required_tables_count,
        'docx_tables': docx_tables,
        'docx_images': docx_images,
        'leftover_figure_placeholders': len(leftover_fig),
        'leftover_table_placeholders': len(leftover_tbl),
    }

    return result


# ==================== Report Generation ====================

def generate_report(results: dict) -> str:
    """生成验证报告"""
    # 总体结论
    statuses = [v.get('status', 'pass') for v in results.values()]
    if 'block' in statuses:
        overall = 'FAIL'
    elif 'warn' in statuses:
        overall = 'WARN'
    else:
        overall = 'PASS'

    lines = ['# DOCX 验证报告\n']
    lines.append(f'## 总体结论\n\n**{overall}**\n')

    # BLOCKER
    blockers = []
    for key, data in results.items():
        if data.get('status') == 'block':
            for issue in data.get('issues', []):
                blockers.append(f'- [{key}] {issue}')
    if blockers:
        lines.append('## BLOCKER\n')
        lines.extend(blockers)
        lines.append('')

    # MAJOR (warn)
    warns = []
    for key, data in results.items():
        if data.get('status') == 'warn':
            for issue in data.get('issues', []):
                warns.append(f'- [{key}] {issue}')
    if warns:
        lines.append('## MAJOR\n')
        lines.extend(warns)
        lines.append('')

    # 检查统计
    lines.append('## 检查统计\n')
    formula = results.get('formula_integrity', {})
    ft = results.get('figure_table', {})
    table = results.get('table_structure', {})
    cite = results.get('citations', {})
    font = results.get('font_paragraph', {})
    open_ = results.get('openability', {})

    lines.append(f'- Markdown 行内公式数: {formula.get("md_inline", "N/A")}')
    lines.append(f'- Markdown 独立公式数: {formula.get("md_block", "N/A")}')
    lines.append(f'- Word oMath 数: {formula.get("docx_omath", "N/A")}')
    lines.append(f'- Word oMathPara 数: {formula.get("docx_omathpara", "N/A")}')
    lines.append(f'- Markdown 图片数: {ft.get("md_images", "N/A")}')
    lines.append(f'- Word 嵌入图片数: {ft.get("docx_images", "N/A")}')
    lines.append(f'- 表格数: {table.get("table_count", "N/A")}')
    lines.append(f'- 表编号: {ft.get("table_nums", [])}')
    lines.append(f'- 图编号: {ft.get("figure_nums", [])}')
    lines.append(f'- 正文引用编号: {cite.get("body_citations", [])}')
    lines.append(f'- 参考文献编号: {cite.get("reference_citations", [])}')
    lines.append(f'- LaTeX 残留: {formula.get("latex_residue", [])}')
    lines.append(f'- 变量缺失句式: {formula.get("loss_sentences", [])}')
    lines.append(f'- 可打开性: {open_.get("status", "N/A")}')

    bold = results.get('long_bold_run', {})
    center = results.get('center_misjudgment', {})
    if bold.get('issues'):
        lines.append(f'- 长加粗段落: {len(bold["issues"])} 处')
    if center.get('issues'):
        lines.append(f'- 居中误判: {len(center["issues"])} 处')

    if table.get('table_details'):
        lines.append(f'- 表格详情:')
        for detail in table['table_details']:
            lines.append(f'  - {detail}')

    # 资产完整性
    assets = results.get('assets_completeness', {})
    if assets.get('stats') and not assets['stats'].get('skipped'):
        lines.append(f'- 资产清单: `{assets["stats"].get("assets_manifest", "N/A")}`')
        lines.append(f'- required 图片: {assets["stats"].get("required_figures", 0)}')
        lines.append(f'- required 表格: {assets["stats"].get("required_tables", 0)}')
        lines.append(f'- docx 嵌入图片: {assets["stats"].get("docx_images", 0)}')
        leftover_fig = assets["stats"].get("leftover_figure_placeholders", 0)
        leftover_tbl = assets["stats"].get("leftover_table_placeholders", 0)
        if leftover_fig or leftover_tbl:
            lines.append(f'- 未替换占位符: [FIGURE:] {leftover_fig} 个, [TABLE:] {leftover_tbl} 个')

    lines.append('')

    # 公式验证详情
    if formula:
        lines.append('## 公式验证\n')
        lines.append(f'- Markdown 行内公式数: {formula.get("md_inline", 0)}')
        lines.append(f'- Markdown 独立公式数: {formula.get("md_block", 0)}')
        lines.append(f'- Word oMath 数: {formula.get("docx_omath", 0)}')
        lines.append(f'- Word oMathPara 数: {formula.get("docx_omathpara", 0)}')
        lines.append(f'- LaTeX 残留: {formula.get("latex_residue", []) or "无"}')
        lines.append(f'- 变量缺失句式: {formula.get("loss_sentences", []) or "无"}')
        lines.append(f'- 判定: {formula.get("status", "unknown").upper()}')
        lines.append('')

    # 建议修复
    if blockers:
        lines.append('## 建议修复\n')
        for b in blockers:
            lines.append(f'1. 修复 {b}')
        lines.append('')

    return '\n'.join(lines)


def check_template_compliance(docx_path: str, template_rules_path: str) -> dict:
    """检查 docx 是否符合模板规则"""
    result = {'status': 'pass', 'issues': [], 'details': []}

    if not template_rules_path or not os.path.exists(template_rules_path):
        return result

    try:
        import json as _json
        rules_data = _json.loads(open(template_rules_path, encoding='utf-8').read())
    except Exception:
        result['issues'].append(f'无法读取 template_rules.json: {template_rules_path}')
        return result

    rules = rules_data.get('rules', {})
    if not rules:
        result['issues'].append('template_rules.json 中无规则定义')
        return result

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(docx_path)
    except Exception:
        result['issues'].append('无法读取 docx，跳过模板合规检查')
        return result

    # 1. 检查标题层级不能全是 Normal
    heading_count = 0
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith('Heading'):
            heading_count += 1
    if heading_count == 0:
        result['status'] = 'block'
        result['issues'].append('BLOCKER: 标题层级全部为 Normal，未使用 Heading 样式')

    # 2. 抽样检查字体是否符合模板规则
    body_rule = rules.get('body', {})
    expected_ea = body_rule.get('font_name_east_asia')
    if expected_ea and body_rule.get('source') == 'explicit_text':
        # 只在模板有明确文字说明时检查
        mismatch_count = 0
        sample_count = 0
        for p in doc.paragraphs[:30]:
            style_name = p.style.name if p.style else ''
            if style_name.startswith('Heading'):
                continue
            for run in p.runs:
                rPr = run._r.find(qn('w:rPr'))
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        ea = rFonts.get(qn('w:eastAsia'))
                        sample_count += 1
                        if ea and ea != expected_ea:
                            mismatch_count += 1
        if sample_count > 0 and mismatch_count > sample_count * 0.5:
            result['status'] = 'warn'
            result['issues'].append(
                f'WARN: 正文中 {mismatch_count}/{sample_count} 个 run 的 eastAsia 字体 '
                f'与模板要求的 "{expected_ea}" 不一致'
            )

    # 3. 记录详情
    for role, rule in rules.items():
        source = rule.get('source', 'unknown')
        if source == 'explicit_text':
            result['details'].append(f'{role}: 使用模板明确说明 ({rule.get("raw_description", "")})')
        elif source == 'word_style':
            result['details'].append(f'{role}: 使用 Word 样式属性')
        else:
            result['details'].append(f'{role}: 使用默认值')

    return result


def main():
    parser = argparse.ArgumentParser(description='DOCX 验证脚本')
    parser.add_argument('--docx', required=True, help='待验证的 docx 文件路径')
    parser.add_argument('--markdown', help='对应的 Markdown 源文件路径')
    parser.add_argument('--tables', help='表格目录路径（可选）')
    parser.add_argument('--figures', help='图片目录路径（可选）')
    parser.add_argument('--template-rules', default=None, help='template_rules.json 路径（可选）')
    parser.add_argument('--assets-manifest', default=None, help='assets_manifest.json 路径（可选）')
    parser.add_argument('--output', required=True, help='验证报告输出路径')
    parser.add_argument('--json-output', help='结构化质量结果 JSON 路径')

    args = parser.parse_args()

    if not os.path.exists(args.docx):
        logger.error('docx 文件不存在: %s', args.docx)
        sys.exit(2)

    results = {}

    # Check A
    results['openability'] = check_openability(args.docx)
    if results['openability']['status'] == 'block':
        # 无法打开则跳过后续检查
        report = generate_report(results)
        _write_report(report, args.output)
        quality = make_quality_result('validate_docx', results, metadata={'docx': args.docx})
        write_json(quality, args.json_output or os.path.splitext(args.output)[0] + '.json')
        sys.exit(2)

    # Check B
    results['formula_integrity'] = check_formula_integrity(args.docx, args.markdown)

    # Check C
    results['figure_table'] = check_figure_table_integrity(
        args.docx, args.markdown, args.tables, args.figures
    )

    # Check D
    results['table_structure'] = check_table_structure(args.docx)

    # Check E
    results['font_paragraph'] = check_font_paragraph(args.docx)

    # Check F
    results['citations'] = check_citations_references(args.docx)

    # Check G
    results['long_bold_run'] = check_long_bold_run(args.docx)

    # Check H
    results['center_misjudgment'] = check_center_misjudgment(args.docx)

    # Check I: Template compliance
    results['template_compliance'] = check_template_compliance(args.docx, args.template_rules)

    # Check J: Assets completeness
    results['assets_completeness'] = check_assets_completeness(args.docx, args.markdown, args.assets_manifest)

    # 生成报告
    report = generate_report(results)
    _write_report(report, args.output)
    quality = make_quality_result('validate_docx', results, metadata={'docx': args.docx})
    json_output = args.json_output or os.path.splitext(args.output)[0] + '.json'
    write_json(quality, json_output)

    # 打印报告
    print(report)
    print(f'\n报告已写入: {args.output}')
    print(f'结构化结果已写入: {json_output}')

    # Exit code
    sys.exit(2 if quality['verdict'] in {'FAIL', 'INCOMPLETE'} else 0)


def _write_report(report: str, output_path: str):
    """写入报告文件"""
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(report)


if __name__ == '__main__':
    main()
