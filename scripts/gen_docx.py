#!/usr/bin/env python3
"""
gen_docx.py -- Markdown draft (with LaTeX math) -> formatted Word document

Pipeline:
    paper_draft.md -> pandoc -> raw.docx -> python-docx post-processing -> paper_final.docx

Usage:
    python gen_docx.py \
        --manifest paper_workspace/00_intake/output/manifest.json \
        --markdown paper_workspace/04_writer/output/paper_draft.md \
        --output paper_workspace/final_paper/paper_final.docx \
        --reference-doc template.docx \
        --tables paper_workspace/03_coder/output/tables \
        --figures paper_workspace/03_coder/output/figures \
        --log paper_workspace/final_paper/docx_build_log.md
"""

import argparse
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import count_md_formulas_from_file as count_markdown_formulas, is_caption

from docx_gen.styles import (
    normalize_heading_styles, normalize_paragraphs,
    normalize_abstract_keywords, set_run_fonts,
    superscript_numeric_citations, _ensure_style,
    _get_body_fonts, _get_title_fonts, _get_heading_fonts,
)
from docx_gen.tables import (
    apply_three_line_tables, fix_table_formatting,
    apply_three_line_table, _get_or_create_tc_borders, _set_tc_border,
)
from docx_gen.assets import (
    render_table_asset_as_markdown, _check_caption_adjacency,
    replace_asset_placeholders, _validate_assets_embedded, insert_images,
)
from docx_gen.formulas import protect_math_objects
from docx_gen.output import (
    save_and_roundtrip_check, write_build_log,
    _make_text_run, _make_superscript_run,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# load_manifest / load_template_rules
# ---------------------------------------------------------------------------

def load_manifest(path: str) -> dict:
    """Load manifest.json; returns empty dict on failure."""
    p = Path(path)
    if not p.exists():
        logger.warning("manifest not found: %s", path)
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as exc:
        logger.warning("manifest parse error: %s", exc)
        return {}


def load_template_rules(path: str | None, manifest: dict = None) -> dict:
    """Load template_rules.json; returns empty dict on failure.

    Resolution order:
    1. Explicit --template-rules path
    2. manifest["template_rules_file"]
    """
    if not path and manifest:
        path = manifest.get("template_rules_file")
    if not path:
        return {}
    p = Path(path)
    if not p.exists():
        logger.warning("template_rules not found: %s", path)
        return {}
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
        logger.info("template_rules loaded: %s (priority: %s)", path, data.get("priority", "unknown"))
        return data
    except (json.JSONDecodeError, OSError) as exc:
        logger.warning("template_rules parse error: %s", exc)
        return {}


def load_assets_manifest(path: str | None, manifest: dict = None) -> dict:
    """Load assets_manifest.json; returns empty dict on failure.

    Resolution order:
    1. Explicit --assets-manifest path
    2. manifest["assets_manifest_file"]
    """
    if not path and manifest:
        path = manifest.get("assets_manifest_file")
    if not path:
        return {}
    p = Path(path)
    if not p.exists():
        logger.warning("assets_manifest not found: %s", path)
        return {}
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
        fig_count = len(data.get("figures", []))
        tbl_count = len(data.get("tables", []))
        logger.info("assets_manifest loaded: %s (%d figures, %d tables)", path, fig_count, tbl_count)
        return data
    except (json.JSONDecodeError, OSError) as exc:
        logger.warning("assets_manifest parse error: %s", exc)
        return {}


# ---------------------------------------------------------------------------
# find_pandoc
# ---------------------------------------------------------------------------

def find_pandoc(cli_path: str | None = None) -> str:
    """Locate pandoc binary with this priority:
    1. cli_path (--pandoc argument)
    2. PANDOC_PATH env var
    3. shutil.which('pandoc')
    4. pypandoc.get_pandoc_path()
    5. Raise RuntimeError
    """
    # 1. CLI argument
    if cli_path:
        if os.path.isfile(cli_path) and os.access(cli_path, os.X_OK):
            return cli_path
        raise RuntimeError(f"--pandoc path not executable: {cli_path}")

    # 2. Environment variable
    env_path = os.environ.get("PANDOC_PATH")
    if env_path and os.path.isfile(env_path) and os.access(env_path, os.X_OK):
        return env_path

    # 3. shutil.which
    which_path = shutil.which("pandoc")
    if which_path:
        return which_path

    # 4. pypandoc
    try:
        import pypandoc
        return pypandoc.get_pandoc_path()
    except (ImportError, OSError):
        pass

    # 5. Not found
    raise RuntimeError(
        "pandoc not found. Install one of the following ways:\n"
        "  - Ubuntu/Debian: sudo apt install pandoc\n"
        "  - macOS: brew install pandoc\n"
        "  - conda: conda install -c conda-forge pandoc\n"
        "  - Or set PANDOC_PATH=/path/to/pandoc\n"
        "  - Or pass --pandoc /path/to/pandoc"
    )


# ---------------------------------------------------------------------------
# run_pandoc
# ---------------------------------------------------------------------------

def run_pandoc(
    markdown: str,
    raw_docx: str,
    reference_doc: str | None,
    pandoc_path: str,
    cwd: str | None = None,
    resource_path: list[str] | None = None,
) -> str:
    """Run pandoc to convert markdown -> docx. Returns pandoc version string.

    Args:
        cwd: Working directory for pandoc (typically markdown's parent dir).
        resource_path: List of directories for pandoc to search for images/resources.
    """
    # Get pandoc version
    version_out = subprocess.run(
        [pandoc_path, "--version"],
        capture_output=True, text=True, timeout=30,
    )
    version_str = version_out.stdout.strip().split("\n")[0] if version_out.stdout else "unknown"

    cmd = [
        pandoc_path,
        markdown,
        "--from", "markdown+tex_math_dollars+tex_math_single_backslash",
        "--to", "docx",
    ]
    if reference_doc and os.path.isfile(reference_doc):
        cmd += ["--reference-doc", reference_doc]
    if resource_path:
        cmd += ["--resource-path", os.pathsep.join(resource_path)]
    cmd += ["-o", raw_docx]

    logger.info("pandoc cmd: %s", " ".join(cmd))
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, cwd=cwd)
    if result.returncode != 0:
        raise RuntimeError(
            f"pandoc failed (exit {result.returncode}):\n"
            f"  stdout: {result.stdout}\n"
            f"  stderr: {result.stderr}"
        )
    if result.stderr:
        logger.warning("pandoc stderr: %s", result.stderr.strip())

    return version_str


# ---------------------------------------------------------------------------
# open_docx_safely
# ---------------------------------------------------------------------------

def open_docx_safely(path: str) -> Document:
    """Open a docx file, raising clear error if missing."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"docx not found: {path}")
    return Document(str(p))


# ---------------------------------------------------------------------------
# Post-processing: center_captions
# ---------------------------------------------------------------------------

def center_captions(doc: Document) -> int:
    """Center paragraphs starting with '表X' or '图X' (table/figure captions).
    Also sets keep_with_next so captions stick to the following table/figure."""
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if is_caption(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = para.paragraph_format
            pf.keep_with_next = True
            # Set fonts on caption runs
            for run in para.runs:
                set_run_fonts(run)
            count += 1
    return count


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def build_docx(args: argparse.Namespace) -> int:
    """Run the full pipeline. Returns 0 on success, 1 on error."""
    warnings = []

    # 0. Load manifest (optional)
    manifest = load_manifest(args.manifest) if args.manifest else {}
    if not manifest and args.manifest:
        warnings.append(f"manifest not found or invalid: {args.manifest}")
    elif not args.manifest:
        warnings.append("manifest: not provided")

    # 0b. Load template rules (optional)
    template_rules = load_template_rules(args.template_rules, manifest)
    if template_rules:
        body_ea, body_ascii = _get_body_fonts(template_rules)
        logger.info("template rules loaded: body fonts = %s / %s", body_ea, body_ascii)
    else:
        body_ea, body_ascii = "宋体", "Times New Roman"

    # 0b2. Template rules mandatory when Word template exists
    word_template_file = manifest.get("word_template_file")
    has_word_template = bool(args.reference_doc or word_template_file)
    if has_word_template and not template_rules:
        msg = (
            "FATAL: Word 模板已指定但 template_rules.json 缺失。"
            "请先运行 extract_word_template_rules.py 生成 template_rules.json。"
        )
        logger.error(msg)
        return 1

    # 0c. Load assets manifest (optional)
    assets = load_assets_manifest(args.assets_manifest, manifest)

    # 1. Find pandoc
    try:
        pandoc_path = find_pandoc(args.pandoc)
    except RuntimeError as exc:
        logger.error(str(exc))
        return 1

    pandoc_source = "cli"
    if args.pandoc:
        pandoc_source = "--pandoc"
    elif os.environ.get("PANDOC_PATH"):
        pandoc_source = "PANDOC_PATH env"
    elif shutil.which("pandoc"):
        pandoc_source = "shutil.which"
    else:
        try:
            import pypandoc  # noqa: F401
            pandoc_source = "pypandoc"
        except ImportError:
            pandoc_source = "unknown"

    # 2. Validate inputs
    md_path = Path(args.markdown)
    if not md_path.exists():
        logger.error("markdown not found: %s", args.markdown)
        return 1

    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    reference_doc = args.reference_doc if args.reference_doc else None
    if reference_doc and not Path(reference_doc).exists():
        logger.warning("reference-doc not found: %s, proceeding without it", reference_doc)
        warnings.append(f"reference-doc not found: {reference_doc}")
        reference_doc = None

    # 3. Count markdown formulas
    inline_count, block_count = count_markdown_formulas(str(md_path))

    # 4. Run pandoc — all processing inside tempdir lifetime
    img_count = 0
    post_steps = {}
    stats = {}

    with tempfile.TemporaryDirectory(prefix="gen_docx_") as tmpdir:
        raw_docx = os.path.join(tmpdir, "raw.docx")

        # 3b. Replace asset placeholders before pandoc
        md_text = md_path.read_text(encoding="utf-8")
        placeholder_stats = {}
        log_entries = []
        if assets:
            # Validate caption adjacency before replacing placeholders
            caption_errors = _check_caption_adjacency(md_text)
            if caption_errors:
                for err in caption_errors:
                    logger.error("CAPTION ADJACENCY: %s", err)
                    log_entries.append(f"CAPTION ADJACENCY ERROR: {err}")
                logger.error("Caption adjacency check failed (%d errors). Aborting.", len(caption_errors))
                _write_build_log(log_entries, args.log)
                return 1

            md_text, placeholder_stats = replace_asset_placeholders(md_text, assets)
            logger.info(
                "asset placeholders: fig=%d replaced/%d missing, tbl=%d replaced/%d missing",
                placeholder_stats.get("figures_replaced", 0),
                len(placeholder_stats.get("figures_missing", [])),
                placeholder_stats.get("tables_replaced", 0),
                len(placeholder_stats.get("tables_missing", [])),
            )

        # Write modified markdown to temp file for pandoc
        tmp_md = os.path.join(tmpdir, "paper_draft_assets.md")
        Path(tmp_md).write_text(md_text, encoding="utf-8")

        # Set up cwd and resource_path for reliable image discovery
        pandoc_cwd = str(md_path.parent)
        resource_dirs = [str(md_path.parent)]
        if args.figures and os.path.isdir(args.figures):
            resource_dirs.append(str(Path(args.figures).resolve()))
            # Also add figures subdirectories
            for sub in Path(args.figures).iterdir():
                if sub.is_dir():
                    resource_dirs.append(str(sub.resolve()))
        if args.tables and os.path.isdir(args.tables):
            resource_dirs.append(str(Path(args.tables).resolve()))

        try:
            pandoc_version = run_pandoc(
                tmp_md, raw_docx, reference_doc, pandoc_path,
                cwd=pandoc_cwd, resource_path=resource_dirs,
            )
        except RuntimeError as exc:
            logger.error("pandoc failed: %s", exc)
            return 1

        # 5. Open raw docx
        doc = open_docx_safely(raw_docx)

        # 6. Post-processing pipeline (all inside tempdir)
        post_steps = {}

        # 6a. Protect math objects (count only, no modification)
        math_count = protect_math_objects(doc)
        post_steps["protect_math_objects"] = f"{math_count} math elements preserved"

        # 6b. Normalize heading styles
        heading_count = normalize_heading_styles(doc)
        post_steps["normalize_heading_styles"] = f"{heading_count} headings adjusted"

        # 6c. Set fonts on all runs in the document body
        # Use template rules for heading-specific fonts; body uses body_ea/body_ascii
        _title_font_done = False
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            # Determine fonts for this paragraph
            ea, ascii = body_ea, body_ascii
            # Paper title: first Normal paragraph with center alignment
            if (
                not _title_font_done
                and style_name == "Normal"
                and para.alignment == WD_ALIGN_PARAGRAPH.CENTER
                and para.text.strip()
            ):
                title_fonts = _get_title_fonts(template_rules)
                if title_fonts:
                    ea, ascii = title_fonts
                # Get title font size from template rules
                _title_size = None
                _tr = template_rules.get("rules", {}).get("title", {})
                if _tr.get("font_size_pt"):
                    _title_size = _tr["font_size_pt"]
                for run in para.runs:
                    set_run_fonts(run, east_asia=ea, ascii_font=ascii, font_size_pt=_title_size)
                _title_font_done = True
                continue
            elif style_name.startswith("Heading 1"):
                hfonts = _get_heading_fonts(template_rules, 1)
                if hfonts:
                    ea, ascii = hfonts
            elif style_name.startswith("Heading 2"):
                hfonts = _get_heading_fonts(template_rules, 2)
                if hfonts:
                    ea, ascii = hfonts
            elif style_name.startswith("Heading 3"):
                hfonts = _get_heading_fonts(template_rules, 3)
                if hfonts:
                    ea, ascii = hfonts
            for run in para.runs:
                set_run_fonts(run, east_asia=ea, ascii_font=ascii)
        post_steps["set_run_fonts"] = f"all runs: eastAsia={body_ea}, ascii={body_ascii}"

        # 6d. Normalize paragraphs (first-line indent)
        para_count = normalize_paragraphs(doc, template_rules)
        post_steps["normalize_paragraphs"] = f"{para_count} paragraphs indented"

        # 6e. Abstract/keywords formatting
        abstract_count = normalize_abstract_keywords(doc)
        post_steps["normalize_abstract_keywords"] = f"{abstract_count} labels formatted"

        # 6f. Superscript citations
        cite_count = superscript_numeric_citations(doc)
        post_steps["superscript_numeric_citations"] = f"{cite_count} citations superscripted"

        # 6g. Center captions
        caption_count = center_captions(doc)
        post_steps["center_captions"] = f"{caption_count} captions centered"

        # 6h. Three-line tables
        table_count = apply_three_line_tables(doc)
        post_steps["apply_three_line_tables"] = f"{table_count} tables styled"

        # 6h2. Table formatting (alignment, font size, cell alignment)
        fmt_count = fix_table_formatting(doc)
        post_steps["fix_table_formatting"] = f"{fmt_count} tables formatted"

        # 6i. Insert images
        img_count = 0
        if args.figures and os.path.isdir(args.figures):
            img_count = insert_images(doc, args.figures)
            post_steps["postprocess_images_inserted"] = f"{img_count} images inserted by postprocess"
        else:
            post_steps["postprocess_images_inserted"] = "skipped (no figures dir)"

        # 7. Save and verify (still inside tempdir)
        stats = save_and_roundtrip_check(doc, str(out_path))

    # with exited — tempdir cleaned up, but output already saved

    # 8. Count markdown images and docx embedded images
    md_text_for_images = md_path.read_text(encoding="utf-8") if md_path.exists() else ""
    md_image_count = len(re.findall(r'!\[.*?\]\(([^)]+)\)', md_text_for_images))
    docx_embedded_images = 0
    if out_path.exists():
        try:
            import zipfile as _zf
            with _zf.ZipFile(str(out_path), 'r') as _z:
                docx_embedded_images = len([n for n in _z.namelist() if n.startswith("word/media/")])
        except Exception:
            pass

    # 8b. Validate required assets
    asset_warnings = _validate_assets_embedded(md_text, assets, placeholder_stats)
    fatal_asset_warnings = asset_warnings  # all asset warnings are fatal
    warnings.extend(asset_warnings)

    # Asset stats for build log
    asset_fig_required = sum(1 for f in assets.get("figures", []) if f.get("required", True)) if assets else 0
    asset_fig_replaced = placeholder_stats.get("figures_replaced", 0)
    asset_tbl_required = sum(1 for t in assets.get("tables", []) if t.get("required", True)) if assets else 0
    asset_tbl_replaced = placeholder_stats.get("tables_replaced", 0)
    asset_missing = len(placeholder_stats.get("figures_missing", [])) + len(placeholder_stats.get("tables_missing", []))

    # 9. Write build log
    log_path = args.log if args.log else str(out_path.parent / "docx_build_log.md")
    write_build_log(
        log_path,
        pandoc_source=pandoc_source,
        pandoc_version=pandoc_version,
        markdown=args.markdown,
        reference_doc=reference_doc or "(none)",
        template_rules_path=args.template_rules or "(none)",
        tables_dir=args.tables or "(none)",
        figures_dir=args.figures or "(none)",
        inline_formulas=inline_count,
        block_formulas=block_count,
        math_omathpara=stats.get("math_omathpara", 0),
        math_omath=stats.get("math_omath", 0),
        paragraphs=stats.get("paragraphs", 0),
        tables=stats.get("tables", 0),
        sections=stats.get("sections", 0),
        markdown_images=md_image_count,
        docx_embedded_images=docx_embedded_images,
        postprocess_images_inserted=img_count,
        post_steps=post_steps,
        warnings=warnings,
        manifest_path=args.manifest or "(not provided)",
        manifest_output_format=manifest.get("output_format", "(unknown)"),
        manifest_pandoc_status=manifest.get("dependency_status", {}).get("pandoc", "(unknown)"),
        assets_manifest_path=args.assets_manifest or "(not provided)",
        asset_fig_required=asset_fig_required,
        asset_fig_replaced=asset_fig_replaced,
        asset_tbl_required=asset_tbl_required,
        asset_tbl_replaced=asset_tbl_replaced,
        asset_missing=asset_missing,
    )

    print(f"OK: {out_path}")
    print(f"  paragraphs: {stats.get('paragraphs', 0)}")
    print(f"  tables: {stats.get('tables', 0)}")
    print(f"  math objects: {stats.get('math_omath', 0)}")
    print(f"  images embedded: {docx_embedded_images}")
    print(f"  images inserted by postprocess: {img_count}")
    print(f"  log: {log_path}")

    if fatal_asset_warnings:
        logger.critical("%d asset issue(s) detected:", len(fatal_asset_warnings))
        for w in fatal_asset_warnings:
            logger.critical("  - %s", w)
        return 1

    return 0


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown draft to formatted Word document via pandoc + python-docx",
    )
    parser.add_argument(
        "--manifest", required=False, default=None,
        help="Path to manifest.json from Stage 0 (optional)",
    )
    parser.add_argument(
        "--markdown", required=True,
        help="Path to paper_draft.md",
    )
    parser.add_argument(
        "--output", required=True,
        help="Output path for paper_final.docx",
    )
    parser.add_argument(
        "--reference-doc", default=None,
        help="Optional Word template for pandoc --reference-doc",
    )
    parser.add_argument(
        "--template-rules", default=None,
        help="Path to template_rules.json (from extract_word_template_rules.py)",
    )
    parser.add_argument(
        "--assets-manifest", default=None,
        help="Path to assets_manifest.json from Stage 3 coder",
    )
    parser.add_argument(
        "--tables", default=None,
        help="Path to tables directory (for reference)",
    )
    parser.add_argument(
        "--figures", default=None,
        help="Path to figures directory for image insertion",
    )
    parser.add_argument(
        "--log", default=None,
        help="Path for build log markdown",
    )
    parser.add_argument(
        "--pandoc", default=None,
        help="Path to pandoc binary (overrides auto-detection)",
    )
    parser.add_argument(
        "-v", "--verbose", action="store_true",
        help="Enable verbose logging",
    )

    args = parser.parse_args()

    # Setup logging
    log_level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(
        level=log_level,
        format="%(levelname)s: %(message)s",
    )

    try:
        rc = build_docx(args)
    except Exception as exc:
        logger.error(str(exc))
        logger.exception("unhandled exception")
        sys.exit(1)

    sys.exit(rc)


if __name__ == "__main__":
    main()
