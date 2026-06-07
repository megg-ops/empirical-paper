#!/usr/bin/env python3
"""
extract_word_template_rules.py -- 从 Word 模板中提取格式规则

读取 Word 模板的完整文本和样式信息，输出两份文件：
  - template_text.md  — 模板中所有文字说明
  - template_rules.json — 综合样式属性和文字说明的结构化规则

用法：
    python scripts/extract_word_template_rules.py \
      --template <template.docx> \
      --output <workspace>/00_intake/output/template_rules.json \
      --text-output <workspace>/00_intake/output/template_text.md

规则优先级：
    模板中的明确文字说明 > Word 样式属性 > fallback 默认值
"""

import argparse
import json
import logging
import os
import re
import sys
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    logger.error("python-docx is required. Install with: pip install python-docx")
    sys.exit(2)

# 确保 docx_gen 模块可导入（scripts/ 目录）
_scripts_dir = Path(__file__).resolve().parent
if str(_scripts_dir) not in sys.path:
    sys.path.insert(0, str(_scripts_dir))

from docx_gen.defaults import FALLBACK_DEFAULTS, PT_TO_CHINESE, CHINESE_TO_PT


# ---------------------------------------------------------------------------
# Font size mapping (Chinese standard) — see docx_gen.defaults for definitions
# ---------------------------------------------------------------------------


def _pt_to_chinese(pt_val):
    """将 pt 值转为中文字号（取最近匹配）"""
    if pt_val is None:
        return None
    # 精确匹配
    if pt_val in PT_TO_CHINESE:
        return PT_TO_CHINESE[pt_val]
    # 最近匹配
    closest = min(PT_TO_CHINESE.keys(), key=lambda x: abs(x - pt_val))
    if abs(closest - pt_val) <= 0.5:
        return PT_TO_CHINESE[closest]
    return f"{pt_val}pt"


def _alignment_to_chinese(align_val):
    """将对齐方式转为中文"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
        WD_ALIGN_PARAGRAPH.CENTER: "居中",
        WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    }
    return mapping.get(align_val, "未知")


# ---------------------------------------------------------------------------
# Extract all text from template
# ---------------------------------------------------------------------------

def extract_template_text(doc: Document) -> str:
    """提取模板中的所有文字内容，包括段落和表格"""
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = para.style.name if para.style else "Normal"
            lines.append(f"[{style_name}] {text}")

    # 表格中的文字
    for i, table in enumerate(doc.tables):
        lines.append(f"\n--- 表格 {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Extract style properties from Word styles
# ---------------------------------------------------------------------------

def extract_style_properties(doc: Document) -> dict:
    """从 Word 样式中提取格式属性"""
    styles = {}

    # 定义要提取的样式名称映射
    style_map = {
        "title": ["Title", "标题"],
        "heading1": ["Heading 1", "标题 1", "heading 1"],
        "heading2": ["Heading 2", "标题 2", "heading 2"],
        "heading3": ["Heading 3", "标题 3", "heading 3"],
        "body": ["Normal", "正文", "Body Text"],
        "abstract": ["Abstract", "摘要"],
        "keywords": [],
        "table_caption": [],
        "figure_caption": [],
        "table": [],
        "references": ["Bibliography", "参考文献"],
    }

    for role, candidate_names in style_map.items():
        styles[role] = _extract_one_style(doc, candidate_names, role)

    return styles


def _extract_one_style(doc: Document, candidate_names: list, role: str) -> dict:
    """提取单个样式的属性"""
    result = {
        "font_name": None,
        "font_name_east_asia": None,
        "font_size_pt": None,
        "font_size_chinese": None,
        "bold": None,
        "alignment": None,
        "first_line_indent_chars": None,
        "line_spacing": None,
        "found_style": False,
    }

    # 尝试从 style 定义中提取
    for style in doc.styles:
        style_name = style.name or ""
        if style_name in candidate_names or any(c.lower() in style_name.lower() for c in candidate_names):
            result["found_style"] = True
            if style.font:
                if style.font.size:
                    pt_val = style.font.size.pt
                    result["font_size_pt"] = pt_val
                    result["font_size_chinese"] = _pt_to_chinese(pt_val)
                if style.font.bold is not None:
                    result["bold"] = style.font.bold
                # 字体名
                if style.font.name:
                    result["font_name"] = style.font.name
                rpr = style.element.find(qn("w:rPr"))
                if rpr is not None:
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is not None:
                        ea = rfonts.get(qn("w:eastAsia"))
                        if ea:
                            result["font_name_east_asia"] = ea
            # 段落属性
            if hasattr(style, "paragraph_format") and style.paragraph_format:
                pf = style.paragraph_format
                if pf.alignment is not None:
                    result["alignment"] = _alignment_to_chinese(pf.alignment)
                if pf.first_line_indent is not None:
                    from docx.shared import Pt as SharedPt
                    try:
                        # 假设小四号(12pt)字体，1字符 = 12pt
                        chars = round(pf.first_line_indent.pt / 12, 1)
                        result["first_line_indent_chars"] = chars
                    except Exception:
                        pass
                if pf.line_spacing is not None:
                    result["line_spacing"] = str(pf.line_spacing)
            break

    # 如果样式定义中没有找到，尝试从实际段落中提取
    if not result["found_style"]:
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name in candidate_names or (
                role == "heading1" and style_name.startswith("Heading 1")
            ) or (
                role == "heading2" and style_name.startswith("Heading 2")
            ) or (
                role == "heading3" and style_name.startswith("Heading 3")
            ):
                result["found_style"] = True
                result.update(_extract_para_properties(para))
                break

    return result


def _extract_para_properties(para) -> dict:
    """从实际段落中提取格式属性"""
    result = {}

    # 对齐
    if para.alignment is not None:
        result["alignment"] = _alignment_to_chinese(para.alignment)

    # 段落格式
    pf = para.paragraph_format
    if pf.first_line_indent is not None:
        try:
            chars = round(pf.first_line_indent.pt / 12, 1)
            result["first_line_indent_chars"] = chars
        except Exception:
            pass
    if pf.line_spacing is not None:
        result["line_spacing"] = str(pf.line_spacing)

    # 从 runs 中提取字体信息
    for run in para.runs:
        if run.font.size:
            result["font_size_pt"] = run.font.size.pt
            result["font_size_chinese"] = _pt_to_chinese(run.font.size.pt)
        if run.font.name:
            result["font_name"] = run.font.name
        if run.font.bold is not None:
            result["bold"] = run.font.bold
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                ea = rfonts.get(qn("w:eastAsia"))
                if ea:
                    result["font_name_east_asia"] = ea
        break  # 只取第一个 run

    return result


# ---------------------------------------------------------------------------
# Parse explicit text rules from template content
# ---------------------------------------------------------------------------

def parse_text_rules(template_text: str) -> dict:
    """从模板文字中解析明确的格式说明

    匹配模式如：
    - "一级标题：三号黑体，居中"
    - "正文：小四宋体，首行缩进2字符，1.5倍行距"
    - "Heading 1: 三号黑体, 居中"
    """
    rules = {}

    # 常见的格式说明模式
    # 分隔符支持：冒号（：/ :）、括号（（/））、破折号（——）
    _sep = r"[：:（(——]"
    patterns = [
        # 中文标题模式 — 冒号格式
        (r"(?:一级|1级|第[一二1]级)?标题[：:]\s*(.+)", "heading1"),
        (r"(?:二级|2级|第[二2]级)?标题[：:]\s*(.+)", "heading2"),
        (r"(?:三级|3级|第[三3]级)?标题[：:]\s*(.+)", "heading3"),
        (r"正[文身][：:]\s*(.+)", "body"),
        (r"摘\s*要[：:]\s*(.+)", "abstract"),
        (r"关键词[：:]\s*(.+)", "keywords"),
        (r"(?:图[题注]图说|图[注名]标题)[：:]\s*(.+)", "figure_caption"),
        (r"(?:表[题注]表说|表[注名]标题)[：:]\s*(.+)", "table_caption"),
        (r"表\s*格[：:]\s*(.+)", "table"),
        (r"参考文献[：:]\s*(.+)", "references"),
        (r"(?:论文)?标题[：:]\s*(.+)", "title"),
        # 中文标题模式 — 括号格式（如 "一级标题（四号楷体）"）
        (r"(?:一级|1级|第[一二1]级)?标题[（(]\s*(.+?)[）)]", "heading1"),
        (r"(?:二级|2级|第[二2]级)?标题[（(]\s*(.+?)[）)]", "heading2"),
        (r"(?:三级|3级|第[三3]级)?标题[（(]\s*(.+?)[）)]", "heading3"),
        (r"正[文身][（(]\s*(.+?)[）)]", "body"),
        (r"摘\s*要[（(]\s*(.+?)[）)]", "abstract"),
        (r"关键词[（(]\s*(.+?)[）)]", "keywords"),
        (r"(?:图[题注]图说|图[注名]标题)[（(]\s*(.+?)[）)]", "figure_caption"),
        (r"(?:表[题注]表说|表[注名]标题)[（(]\s*(.+?)[）)]", "table_caption"),
        (r"表\s*格[（(]\s*(.+?)[）)]", "table"),
        (r"参考文献[（(]\s*(.+?)[）)]", "references"),
        (r"(?:论文)?标题[（(]\s*(.+?)[）)]", "title"),
        # 英文标题模式
        (r"Heading\s*1[：:]\s*(.+)", "heading1"),
        (r"Heading\s*2[：:]\s*(.+)", "heading2"),
        (r"Heading\s*3[：:]\s*(.+)", "heading3"),
        (r"Body[：:]\s*(.+)", "body"),
        (r"Abstract[：:]\s*(.+)", "abstract"),
        (r"(?:Table\s*)?Caption[：:]\s*(.+)", "table_caption"),
        (r"(?:Figure\s*)?Caption[：:]\s*(.+)", "figure_caption"),
    ]

    for line in template_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # 去掉样式前缀 [Heading 1] 等
        clean_line = re.sub(r"^\[.*?\]\s*", "", line)

        for pattern, role in patterns:
            m = re.search(pattern, clean_line, re.IGNORECASE)
            if m:
                desc = m.group(1).strip()
                parsed = _parse_format_description(desc)
                if parsed:
                    rules[role] = parsed
                break

    return rules


def _parse_format_description(desc: str) -> dict | None:
    """解析格式描述文字，提取字号、字体、对齐等信息

    示例输入：
    - "三号黑体，居中"
    - "小四宋体，首行缩进2字符，1.5倍行距"
    - "五号Times New Roman，右对齐"
    """
    result = {
        "source": "explicit_text",
        "description": desc,
    }
    found_any = False

    # 提取字号
    for size_name, pt_val in CHINESE_TO_PT.items():
        if size_name in desc:
            result["font_size_chinese"] = size_name
            result["font_size_pt"] = pt_val
            found_any = True
            break

    # 也匹配 "X号" 格式（如 "小四号"）
    m = re.search(r"([小]?[一二三四五六七八]号)", desc)
    if m and "font_size_chinese" not in result:
        size_name = m.group(1)
        if size_name in CHINESE_TO_PT:
            result["font_size_chinese"] = size_name
            result["font_size_pt"] = CHINESE_TO_PT[size_name]
            found_any = True

    # 提取字体
    chinese_fonts = [
        (r"黑体", "黑体"),
        (r"宋体", "宋体"),
        (r"楷体", "楷体"),
        (r"仿宋", "仿宋"),
    ]
    latin_fonts = [
        (r"Times\s*New\s*Roman", "Times New Roman"),
        (r"Arial", "Arial"),
    ]
    for pattern, font_name in chinese_fonts:
        if re.search(pattern, desc, re.IGNORECASE):
            result["font_name"] = font_name
            result["font_name_east_asia"] = font_name
            found_any = True
            break
    if "font_name" not in result:
        for pattern, font_name in latin_fonts:
            if re.search(pattern, desc, re.IGNORECASE):
                result["font_name"] = font_name
                found_any = True
                break

    # 提取对齐方式
    if "居中" in desc:
        result["alignment"] = "居中"
        found_any = True
    elif "左对齐" in desc:
        result["alignment"] = "左对齐"
        found_any = True
    elif "右对齐" in desc:
        result["alignment"] = "右对齐"
        found_any = True
    elif "两端对齐" in desc:
        result["alignment"] = "两端对齐"
        found_any = True

    # 提取首行缩进
    m = re.search(r"首行缩进\s*(\d+)\s*字符", desc)
    if m:
        result["first_line_indent_chars"] = int(m.group(1))
        found_any = True

    # 提取行距
    m = re.search(r"([\d.]+)\s*倍行距", desc)
    if m:
        result["line_spacing"] = f"{m.group(1)}"
        found_any = True

    # 提取加粗
    if "加粗" in desc or "粗体" in desc:
        result["bold"] = True
        found_any = True

    # 三线表
    if "三线表" in desc or "三线" in desc:
        result["table_style"] = "three_line"
        found_any = True

    return result if found_any else None


# ---------------------------------------------------------------------------
# Merge: explicit text rules > style properties > fallback defaults
# ---------------------------------------------------------------------------

# FALLBACK_DEFAULTS imported from docx_gen.defaults


def merge_rules(style_props: dict, text_rules: dict) -> dict:
    """合并规则：explicit_text > style_props > fallback"""
    merged = {}
    warnings = []

    for role, fallback in FALLBACK_DEFAULTS.items():
        merged[role] = {}
        # Start with fallback
        for key, val in fallback.items():
            merged[role][key] = val
        # Override with style properties (if found)
        if role in style_props and style_props[role].get("found_style"):
            for key, val in style_props[role].items():
                if val is not None and key != "found_style":
                    merged[role][key] = val
        # Override with explicit text rules (highest priority)
        if role in text_rules:
            for key, val in text_rules[role].items():
                if val is not None and key not in ("source", "description"):
                    merged[role][key] = val
            merged[role]["source"] = "explicit_text"
            merged[role]["raw_description"] = text_rules[role].get("description", "")
        elif role in style_props and style_props[role].get("found_style"):
            merged[role]["source"] = "word_style"
        else:
            merged[role]["source"] = "fallback_default"
            warnings.append(f"{role}: 未在模板中找到样式或文字说明，使用默认值")

    return merged, warnings


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="从 Word 模板中提取格式规则"
    )
    parser.add_argument(
        "--template", required=True,
        help="Word 模板文件路径 (.docx)",
    )
    parser.add_argument(
        "--output", required=True,
        help="template_rules.json 输出路径",
    )
    parser.add_argument(
        "--text-output", required=True,
        help="template_text.md 输出路径",
    )
    args = parser.parse_args()

    template_path = Path(args.template)
    if not template_path.exists():
        logger.error("模板文件不存在: %s", args.template)
        sys.exit(2)

    # 1. 打开模板
    doc = Document(str(template_path))

    # 2. 提取模板文字
    template_text = extract_template_text(doc)

    # 3. 提取样式属性
    style_props = extract_style_properties(doc)

    # 4. 解析文字中的格式说明
    text_rules = parse_text_rules(template_text)

    # 5. 合并规则
    merged_rules, warnings = merge_rules(style_props, text_rules)

    # 6. 构建输出 JSON
    output = {
        "source_template": str(template_path),
        "template_text_extracted": bool(template_text.strip()),
        "style_properties_extracted": any(
            v.get("found_style") for v in style_props.values()
        ),
        "rules": merged_rules,
        "priority": "explicit_text_rules > word_style_properties > fallback_defaults",
        "warnings": warnings,
    }

    # 7. 写入 template_rules.json
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(output, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    # 8. 写入 template_text.md
    text_output_path = Path(args.text_output)
    text_output_path.parent.mkdir(parents=True, exist_ok=True)
    md_content = f"# 模板文字内容\n\n来源：`{template_path.name}`\n\n"
    md_content += "## 模板正文\n\n```\n"
    md_content += template_text
    md_content += "\n```\n\n"
    if text_rules:
        md_content += "## 检测到的格式说明\n\n"
        for role, rule in text_rules.items():
            md_content += f"- **{role}**: {rule.get('description', '')}\n"
        md_content += "\n"
    md_content += "## 规则优先级\n\n"
    md_content += "模板中的明确文字说明 > Word 样式属性 > skill 默认格式\n"
    text_output_path.write_text(md_content, encoding="utf-8")

    # 9. 输出摘要
    print(f"模板规则已提取:")
    print(f"  - rules: {output_path}")
    print(f"  - text:  {text_output_path}")
    print(f"  - 检测到 {len(text_rules)} 条文字规则")
    print(f"  - {len(warnings)} 条警告")
    if warnings:
        for w in warnings:
            print(f"    ⚠ {w}")


if __name__ == "__main__":
    main()
