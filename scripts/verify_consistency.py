#!/usr/bin/env python3
"""验证论文与结构化分析结果的一致性。

Markdown、LaTeX 与 DOCX 共用同一组判定函数；``results.json`` 是数字真源。
Markdown 报告只供人阅读，最终门禁读取同名 JSON 结果。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from check_word_count import load_requirement
from quality_contract import make_quality_result, write_json
from utils import count_paper_words, evaluate_word_count, extract_md_citation_numbers


NUMBER_RE = re.compile(r"(?<![\w.])[-+]?\d+(?:\.\d+)?%?(?![\w.])")
PLACEHOLDER_RE = re.compile(
    r"TODO|TBD|FIXME|待补|待填写|待完善|待确认|\[\?\]|\{\{[^}]+\}\}", re.IGNORECASE,
)
ASSET_TOKEN_RE = re.compile(r"\[(TABLE|FIGURE):\s*([^\]]+)\]", re.IGNORECASE)


def _read_json(path: str) -> dict:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _paper_text(path: Path, paper_format: str) -> str:
    if paper_format == "docx":
        return _docx_text(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def _asset_items(manifest: dict) -> list[dict]:
    return [*manifest.get("tables", []), *manifest.get("figures", [])]


def _load_assets(results_json: Path) -> tuple[dict, Path | None]:
    path = results_json.with_name("assets_manifest.json")
    if not path.exists():
        return {"tables": [], "figures": []}, None
    return _read_json(str(path)), path


def _resolve_asset_path(raw: str, manifest_path: Path) -> Path:
    candidate = Path(raw)
    if candidate.is_absolute():
        return candidate
    # 先相对 manifest，再兼容以项目工作目录为基准的旧 manifest。
    local = manifest_path.parent / candidate
    return local if local.exists() else candidate


def _expand_structured_assets(text: str, manifest: dict, manifest_path: Path | None) -> str:
    """把 Markdown/LaTeX 中的结构化表格正文加入数字核对语料。"""
    if manifest_path is None:
        return text
    additions = []
    for item in manifest.get("tables", []):
        raw = item.get("path") or item.get("file")
        if not raw:
            continue
        path = _resolve_asset_path(str(raw), manifest_path)
        if path.exists() and path.suffix.lower() in {".md", ".csv", ".tex", ".txt"}:
            additions.append(path.read_text(encoding="utf-8", errors="ignore"))
    return text + ("\n" + "\n".join(additions) if additions else "")


def _contains_form(text: str, form: str) -> bool:
    escaped = re.escape(str(form))
    return re.search(rf"(?<![\w.]){escaped}(?![\w.])", text) is not None


def _check_structured_numbers(text: str, results: dict) -> dict:
    values = results.get("reportable_values")
    if not isinstance(values, list) or not values:
        return {"status": "BLOCKER", "issues": ["results.json.reportable_values 缺失或为空"]}

    issues: list[str] = []
    invalid_entries = 0
    checked = 0
    for entry in values:
        key = entry.get("key", "<unknown>")
        forms = [str(v) for v in entry.get("allowed_text_forms", []) if str(v)]
        if not forms and entry.get("value_display") is not None:
            forms = [str(entry["value_display"])]
        if entry.get("value_raw") is None or not forms:
            invalid_entries += 1
            issues.append(f"{key}: 缺少 value_raw 或 allowed_text_forms")
            continue
        if entry.get("must_report"):
            checked += 1
            if not any(_contains_form(text, form) for form in forms):
                issues.append(f"{key}: 必报值未以允许精度出现在正文或表格中（允许：{', '.join(forms)}）")
    status = "BLOCKER" if issues else "PASS"
    return {
        "status": status,
        "issues": issues,
        "reportable_value_count": len(values),
        "must_report_checked": checked,
        "invalid_entries": invalid_entries,
    }


def _reference_numbers(text: str) -> set[int]:
    marker = re.search(r"(?im)^\s*(?:#{1,6}\s*)?(?:参考文献|references)\s*$", text)
    if not marker:
        return set()
    section = text[marker.end():]
    return {int(value) for value in re.findall(r"(?m)^\s*\[(\d+)\]", section)}


def _check_citations(text: str) -> dict:
    refs = _reference_numbers(text)
    body_match = re.search(r"(?im)^\s*(?:#{1,6}\s*)?(?:参考文献|references)\s*$", text)
    body = text[:body_match.start()] if body_match else text
    cited = set(extract_md_citation_numbers(body))
    issues = []
    if cited and not refs:
        issues.append("正文含数字引用，但未识别到参考文献列表")
    missing = sorted(cited - refs)
    unused = sorted(refs - cited)
    if missing:
        issues.append(f"正文引用缺少文献条目: {missing}")
    noncontinuous = bool(refs and refs != set(range(1, max(refs) + 1)))
    if noncontinuous:
        issues.append("参考文献编号不连续")
    return {
        "status": "BLOCKER" if missing or (cited and not refs) or noncontinuous else "WARN" if unused else "PASS",
        "issues": issues + ([f"未在正文引用的文献编号: {unused}"] if unused else []),
        "cited": sorted(cited),
        "references": sorted(refs),
    }


def _check_assets(text: str, manifest: dict, manifest_path: Path | None, paper_format: str) -> dict:
    declared = {str(item.get("id")) for item in _asset_items(manifest) if item.get("id")}
    tokens = {(kind.upper(), asset_id.strip()) for kind, asset_id in ASSET_TOKEN_RE.findall(text)}
    referenced = {asset_id for _, asset_id in tokens}
    issues: list[str] = []
    missing_files: list[str] = []
    if manifest_path:
        for item in _asset_items(manifest):
            raw = item.get("path") or item.get("file")
            if raw and not _resolve_asset_path(str(raw), manifest_path).exists():
                missing_files.append(str(raw))
    undeclared = sorted(referenced - declared)
    if undeclared:
        issues.append(f"正文引用未声明的资产 ID: {undeclared}")
    if missing_files:
        issues.append(f"assets_manifest 声明的文件不存在: {missing_files}")
    # Markdown/LaTeX 应引用全部声明资产；DOCX 的嵌入完整性由 validate_docx 负责。
    unused = sorted(declared - referenced) if paper_format != "docx" else []
    if unused:
        issues.append(f"已声明但未在正文引用的资产 ID: {unused}")
    return {
        "status": "BLOCKER" if undeclared or missing_files else "WARN" if unused else "PASS",
        "issues": issues,
        "declared_count": len(declared),
        "referenced_count": len(referenced),
    }


def _check_placeholders(text: str) -> dict:
    matches = sorted(set(match.group(0) for match in PLACEHOLDER_RE.finditer(text)))
    return {
        "status": "BLOCKER" if matches else "PASS",
        "issues": [f"存在未解决占位符: {matches}"] if matches else [],
        "matches": matches,
    }


def _check_format_integrity(text: str, paper_format: str) -> dict:
    issues = []
    if paper_format in {"markdown", "docx"}:
        residues = sorted(set(re.findall(r"\\(?:begin|end|includegraphics|input|cite|ref)\b", text)))
        if residues:
            issues.append(f"存在未转换的 LaTeX 命令: {residues}")
    if paper_format == "docx" and ASSET_TOKEN_RE.search(text):
        issues.append("DOCX 中仍存在未替换的 TABLE/FIGURE 占位符")
    return {"status": "BLOCKER" if issues else "PASS", "issues": issues}


def _check_word_count(text: str, manifest: str | None, skipped: bool) -> dict:
    if skipped:
        return {"status": "PASS", "issues": [], "skipped": True, "reason": "Stage 4 已执行独立字数门禁"}
    if not manifest:
        return {"status": "BLOCKER", "issues": ["未提供 --manifest，不能执行用户规则下的字数检查"]}
    try:
        requirement = load_requirement(manifest)
        counts = count_paper_words(text, requirement["scope"])
        evaluation = evaluate_word_count(counts["total"], requirement)
    except (OSError, KeyError, TypeError, ValueError, json.JSONDecodeError) as exc:
        return {"status": "BLOCKER", "issues": [f"字数规则或统计失败: {exc}"]}
    status = "BLOCKER" if evaluation["status"] == "SHORT" else "WARN" if evaluation["status"] == "OVER" else "PASS"
    return {"status": status, "issues": [], **counts, "evaluation": evaluation, "requirement": requirement}


def verify_paper(
    paper_path: str,
    results_json: str,
    *,
    paper_format: str = "auto",
    manifest: str | None = None,
    skip_word_count: bool = False,
) -> dict:
    paper = Path(paper_path)
    if not paper.exists():
        return make_quality_result("verify_consistency", {
            "input": {"status": "BLOCKER", "issues": [f"论文文件不存在: {paper}"]},
        })
    results_path = Path(results_json)
    if not results_path.exists():
        return make_quality_result("verify_consistency", {
            "input": {"status": "BLOCKER", "issues": [f"results.json 不存在: {results_path}"]},
        })
    if paper_format == "auto":
        paper_format = {".md": "markdown", ".docx": "docx", ".tex": "latex"}.get(paper.suffix.lower(), "markdown")
    try:
        text = _paper_text(paper, paper_format)
        results = _read_json(str(results_path))
        assets, assets_path = _load_assets(results_path)
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        return make_quality_result("verify_consistency", {
            "input": {"status": "BLOCKER", "issues": [f"输入读取失败: {exc}"]},
        })

    number_text = _expand_structured_assets(text, assets, assets_path) if paper_format != "docx" else text
    checks = {
        "structured_numbers": _check_structured_numbers(number_text, results),
        "citations": _check_citations(text),
        "assets": _check_assets(text, assets, assets_path, paper_format),
        "placeholders": _check_placeholders(text),
        "format_integrity": _check_format_integrity(text, paper_format),
        "word_count": _check_word_count(text, manifest, skip_word_count),
    }
    return make_quality_result(
        "verify_consistency", checks,
        metadata={"paper": str(paper), "format": paper_format, "results_json": str(results_path)},
    )


def format_report(result: dict) -> str:
    lines = [
        "# 论文一致性检查报告", "", f"- 结构化结论：**{result['verdict']}**", "",
        "| 检查项 | 状态 |", "|---|---|",
    ]
    for name, check in result["checks"].items():
        lines.append(f"| {name} | {check['status']} |")
    lines.extend(["", "## 问题", ""])
    found = False
    for name, check in result["checks"].items():
        for issue in check.get("issues", []):
            found = True
            lines.append(f"- **{check['status']}** `{name}`：{issue}")
    if not found:
        lines.append("- 无")
    lines.extend(["", "> 最终门禁读取同名 JSON；本 Markdown 不作为机器真源。", ""])
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="论文与分析结果一致性检查")
    parser.add_argument("--paper", required=True, help="paper_draft.md/.tex 或 paper_final.docx")
    parser.add_argument("--results", help="兼容的人类可读结果摘要；不作为数字真源")
    parser.add_argument("--results-json", required=True, help="结构化数字真源 results.json")
    parser.add_argument("--output", required=True, help="Markdown 报告")
    parser.add_argument("--json-output", help="结构化报告；默认与 --output 同名 .json")
    parser.add_argument("--format", choices=["auto", "latex", "markdown", "docx"], default="auto")
    parser.add_argument("--manifest", help="未使用 --skip-word-count 时必填")
    parser.add_argument("--skip-word-count", action="store_true")
    args = parser.parse_args()

    result = verify_paper(
        args.paper, args.results_json, paper_format=args.format,
        manifest=args.manifest, skip_word_count=args.skip_word_count,
    )
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(format_report(result), encoding="utf-8")
    json_output = args.json_output or str(output.with_suffix(".json"))
    write_json(result, json_output)
    print(format_report(result))
    print(f"报告已写入: {output}\n结构化结果已写入: {json_output}")
    sys.exit(0 if result["verdict"] in {"PASS", "WARN"} else 2)


if __name__ == "__main__":
    main()
