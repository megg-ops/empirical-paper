"""输出与构建日志 -- 保存文档、roundtrip 检查、构建日志写入。"""

import copy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _make_text_run(text: str, source_run) -> OxmlElement:
    """Create a normal run element copying formatting from source_run."""
    r = OxmlElement("w:r")
    # Copy run properties from source (without superscript)
    rpr_src = source_run._element.find(qn("w:rPr"))
    if rpr_src is not None:
        rpr = copy.deepcopy(rpr_src)
        # Remove any existing vertAlign
        va = rpr.find(qn("w:vertAlign"))
        if va is not None:
            rpr.remove(va)
        r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _make_superscript_run(text: str, source_run) -> OxmlElement:
    """Create a superscript run element."""
    r = OxmlElement("w:r")
    rpr_src = source_run._element.find(qn("w:rPr"))
    if rpr_src is not None:
        rpr = copy.deepcopy(rpr_src)
    else:
        rpr = OxmlElement("w:rPr")
    # Add superscript
    va = rpr.find(qn("w:vertAlign"))
    if va is not None:
        rpr.remove(va)
    va = OxmlElement("w:vertAlign")
    va.set(qn("w:val"), "superscript")
    rpr.append(va)
    # Ensure rFonts is present
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def save_and_roundtrip_check(doc: Document, output_path: str) -> dict:
    """Save document and verify by re-opening. Returns stats dict."""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(out))

    # Roundtrip check: re-open and count elements
    doc2 = Document(str(out))
    stats = {
        "paragraphs": len(doc2.paragraphs),
        "tables": len(doc2.tables),
        "sections": len(doc2.sections),
    }

    # Count math objects
    body = doc2.element.body
    stats["math_omathpara"] = len(body.findall(f".//{qn('m:oMathPara')}"))
    omath_all = body.findall(f".//{qn('m:oMath')}")
    stats["math_omath"] = len(omath_all)

    return stats


def write_build_log(log_path: str, **kwargs) -> None:
    """Write a markdown build log with all pipeline information."""
    lines = [
        "# DOCX Build Log",
        "",
        f"**Generated**: {_now()}",
        "",
        "## Environment",
        "",
        f"- **pandoc source**: {kwargs.get('pandoc_source', 'N/A')}",
        f"- **pandoc version**: {kwargs.get('pandoc_version', 'N/A')}",
        "",
        "## Input",
        "",
        f"- **markdown**: `{kwargs.get('markdown', 'N/A')}`",
        f"- **reference-doc**: `{kwargs.get('reference_doc', 'N/A')}`",
        f"- **tables dir**: `{kwargs.get('tables_dir', 'N/A')}`",
        f"- **figures dir**: `{kwargs.get('figures_dir', 'N/A')}`",
        "",
        "## Formula Counts (Markdown)",
        "",
        f"- **inline ($)**: {kwargs.get('inline_formulas', 0)}",
        f"- **block ($$)**: {kwargs.get('block_formulas', 0)}",
        "",
        "## Formula Counts (Word)",
        "",
        f"- **m:oMathPara (block)**: {kwargs.get('math_omathpara', 0)}",
        f"- **m:oMath (inline + block)**: {kwargs.get('math_omath', 0)}",
        "",
        "## Document Stats",
        "",
        f"- **paragraphs**: {kwargs.get('paragraphs', 0)}",
        f"- **tables**: {kwargs.get('tables', 0)}",
        f"- **sections**: {kwargs.get('sections', 0)}",
        "",
        "## Image Counts",
        "",
        f"- **Markdown images**: {kwargs.get('markdown_images', 0)}",
        f"- **Images embedded in docx**: {kwargs.get('docx_embedded_images', 0)}",
        f"- **Images inserted by postprocess**: {kwargs.get('postprocess_images_inserted', 0)}",
        "",
        "## Asset Counts",
        "",
        f"- **assets_manifest**: `{kwargs.get('assets_manifest_path', '(not provided)')}`",
        f"- **figures required**: {kwargs.get('asset_fig_required', 0)}",
        f"- **figures replaced**: {kwargs.get('asset_fig_replaced', 0)}",
        f"- **tables required**: {kwargs.get('asset_tbl_required', 0)}",
        f"- **tables replaced**: {kwargs.get('asset_tbl_replaced', 0)}",
        f"- **missing assets**: {kwargs.get('asset_missing', 0)}",
        "",
        "## Manifest",
        "",
        f"- **manifest_path**: `{kwargs.get('manifest_path', '(not provided)')}`",
        f"- **template_rules**: `{kwargs.get('template_rules_path', '(not provided)')}`",
        f"- **output_format**: {kwargs.get('manifest_output_format', '(unknown)')}",
        f"- **pandoc dependency**: {kwargs.get('manifest_pandoc_status', '(unknown)')}",
        "",
        "## Post-processing Steps",
        "",
    ]

    steps = kwargs.get("post_steps", {})
    for step_name, detail in steps.items():
        lines.append(f"- **{step_name}**: {detail}")

    warnings = kwargs.get("warnings", [])
    if warnings:
        lines.append("")
        lines.append("## Warnings")
        lines.append("")
        for w in warnings:
            lines.append(f"- {w}")

    lines.append("")

    log_file = Path(log_path)
    log_file.parent.mkdir(parents=True, exist_ok=True)
    log_file.write_text("\n".join(lines), encoding="utf-8")
