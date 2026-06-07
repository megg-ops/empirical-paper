"""表格处理 -- 三线表、表格格式化。"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docx_gen.styles import set_run_fonts


def _get_or_create_tc_borders(cell):
    """Get or create w:tcBorders on a table cell, avoiding duplicates."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    return borders


def _set_tc_border(cell, edge: str, val: str = "single", sz: str = "8"):
    """Set a single border edge on a table cell. Reuses existing element."""
    borders = _get_or_create_tc_borders(cell)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), sz)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")


def apply_three_line_table(table) -> None:
    """Apply three-line style to a single table:
    - Header row: top thick (12), bottom thin (6)
    - Last row: bottom thick (12)
    - All other borders: none
    """
    nrows = len(table.rows)
    if nrows == 0:
        return

    # First: clear all borders on all cells
    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
                _set_tc_border(cell, edge, val="none", sz="0")

    # Header row: top + bottom
    for cell in table.rows[0].cells:
        _set_tc_border(cell, "top", "single", "12")
        _set_tc_border(cell, "bottom", "single", "6")

    # Last row: bottom only
    for cell in table.rows[-1].cells:
        _set_tc_border(cell, "bottom", "single", "12")


def apply_three_line_tables(doc: Document) -> int:
    """Apply three-line style to all tables in document."""
    count = 0
    for table in doc.tables:
        apply_three_line_table(table)
        # Also set fonts in all table cells
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_fonts(run)
        count += 1
    return count


def fix_table_formatting(doc: Document) -> int:
    """Apply table-level formatting: alignment, font size, cell alignment.

    - Center the table itself on the page.
    - Font size: 9pt for normal tables (< 8 cols), 8pt for wide tables (>= 8 cols).
    - Header row (first row): all cells centered.
    - Data rows: first column left-aligned, remaining columns centered.
    Returns count of tables processed.
    """
    count = 0
    for table in doc.tables:
        # --- table-level centering ---
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        ncols = len(table.columns)
        font_size = Pt(8) if ncols >= 8 else Pt(9)
        nrows = len(table.rows)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    # --- font size ---
                    for run in para.runs:
                        run.font.size = font_size
                    # --- cell alignment ---
                    if row_idx == 0:
                        # Header row: all cells centered
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        if col_idx == 0:
                            # First column: left-aligned (default)
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            # Other columns: centered
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        count += 1
    return count
