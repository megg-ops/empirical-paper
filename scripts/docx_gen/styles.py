"""样式处理 -- 标题层级、字体、段落格式、摘要关键词、引用上标。"""

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from utils import is_caption
from docx_gen.output import _make_text_run, _make_superscript_run


def _ensure_style(doc: Document, name: str) -> None:
    """Ensure a built-in heading style exists in the document."""
    try:
        doc.styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def normalize_heading_styles(doc: Document) -> int:
    """Fix heading hierarchy:
    - First Heading 1 -> Normal + centered (paper title)
    - Heading 2 -> Heading 1
    - Heading 3 -> Heading 2
    - Heading 4 -> Heading 3
    Returns count of headings adjusted.
    """
    # Ensure target styles exist before assigning
    for s in ("Heading 1", "Heading 2", "Heading 3"):
        _ensure_style(doc, s)

    count = 0
    first_h1_seen = False

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "Heading 1":
            if not first_h1_seen:
                # Paper title -> Normal, centered
                para.style = doc.styles["Normal"]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_h1_seen = True
                count += 1
            # else: keep as Heading 1 (was already Heading 2 in markdown)
        elif style_name == "Heading 2":
            para.style = doc.styles["Heading 1"]
            count += 1
        elif style_name == "Heading 3":
            para.style = doc.styles["Heading 2"]
            count += 1
        elif style_name == "Heading 4":
            para.style = doc.styles["Heading 3"]
            count += 1

    return count


def _get_body_fonts(template_rules: dict) -> tuple[str, str]:
    """Extract (east_asia, ascii) fonts from template rules for body text.

    Falls back to 宋体 / Times New Roman if not specified.
    """
    rules = template_rules.get("rules", {})
    body = rules.get("body", {})
    east_asia = body.get("font_name_east_asia") or body.get("font_name", "宋体")
    # For ascii, prefer a Latin font name; if body font is Chinese, use default
    ascii_font = body.get("font_name", "Times New Roman")
    # If the font is a Chinese font name, check if there's a separate ascii hint
    if ascii_font in ("宋体", "黑体", "楷体", "仿宋"):
        ascii_font = "Times New Roman"
    return east_asia, ascii_font


def _get_title_fonts(template_rules: dict) -> tuple[str, str] | None:
    """Extract fonts for the paper title from template rules.

    Returns (east_asia, ascii) or None if no title rule found.
    """
    rules = template_rules.get("rules", {})
    title = rules.get("title", {})
    if not title:
        return None
    east_asia = title.get("font_name_east_asia") or title.get("font_name")
    ascii_font = title.get("font_name")
    if not east_asia and not ascii_font:
        return None
    if not east_asia:
        east_asia = "黑体"
    if not ascii_font or ascii_font in ("宋体", "黑体", "楷体", "仿宋"):
        ascii_font = "Times New Roman"
    return east_asia, ascii_font


def _get_heading_fonts(template_rules: dict, heading_level: int) -> tuple[str, str] | None:
    """Extract fonts for a specific heading level from template rules.

    Returns (east_asia, ascii) or None if no heading rule found.
    """
    rules = template_rules.get("rules", {})
    key = f"heading{heading_level}"
    heading = rules.get(key, {})
    if not heading or heading.get("source") == "fallback_default":
        return None
    east_asia = heading.get("font_name_east_asia") or heading.get("font_name")
    ascii_font = heading.get("font_name")
    if not east_asia and not ascii_font:
        return None
    if not east_asia:
        east_asia = "宋体"
    if not ascii_font or ascii_font in ("宋体", "黑体", "楷体", "仿宋"):
        ascii_font = "Times New Roman"
    return east_asia, ascii_font


def set_run_fonts(
    run,
    east_asia: str = "\u5b8b\u4f53",
    ascii_font: str = "Times New Roman",
    font_size_pt: float | None = None,
) -> None:
    """Set font on a single run: east_asia + ascii, optionally font size."""
    rpr = run._element.get_or_add_rPr()

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)

    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)

    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def normalize_paragraphs(doc: Document, template_rules: dict = None) -> int:
    """Set first-line indent of 2 chars (approx 480 twips = 2 * 12pt * 20twips/pt * 2)
    for body paragraphs that are not headings, captions, or special blocks.

    If template_rules provides body.first_line_indent_chars, use that instead.
    """
    # Heading style names to skip
    heading_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"}

    # Determine indent from template rules or fallback to 2 chars
    indent_chars = 2
    if template_rules:
        body_rules = template_rules.get("rules", {}).get("body", {})
        if "first_line_indent_chars" in body_rules:
            indent_chars = body_rules["first_line_indent_chars"]
    INDENT_TWIPS = int(indent_chars * 240)  # 1 char approx 12pt = 240 twips

    count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name in heading_styles:
            continue
        # Skip centered paragraphs (likely titles)
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue
        # Skip empty paragraphs
        text = para.text.strip()
        if not text:
            continue
        # Skip table/figure captions
        if is_caption(text):
            continue
        # Skip if starts with special labels
        if text.startswith("\u6458\u8981") or text.startswith("\u5173\u952e\u8bcd"):
            continue

        # Apply first-line indent
        ppr = para._element.get_or_add_pPr()
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        ind.set(qn("w:firstLineChars"), "200")
        ind.set(qn("w:firstLine"), str(INDENT_TWIPS))

        count += 1

    return count


def normalize_abstract_keywords(doc: Document) -> int:
    """Format '摘要' and '关键词' labels with bold styling.

    Only bold these labels when they appear at the start of the paragraph
    (followed by a colon, dash, or similar separator), to avoid false
    positives when the words appear in body text (e.g. '关键词频次').
    """
    # Labels to bold: must appear at paragraph start followed by separator
    _label_re = re.compile(r"^(摘要|关键词)\s*[：:、\-]")

    count = 0
    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
        # Only process paragraphs that start with a label
        if not _label_re.match(text.lstrip()):
            continue
        # Bold only the label portion (first run(s) up to the separator)
        label_bold_done = False
        for run in para.runs:
            run_text = run.text
            if not label_bold_done and (
                "\u6458\u8981" in run_text or "\u5173\u952e\u8bcd" in run_text
            ):
                run.bold = True
                set_run_fonts(run)
                # If separator is also in this run, stop bolding after this
                if "\uff1a" in run_text or ":" in run_text:
                    label_bold_done = True
                count += 1
            else:
                if not label_bold_done:
                    # Separator might be in a separate run
                    if "\uff1a" in run_text or ":" in run_text:
                        label_bold_done = True
                set_run_fonts(run)
    return count


def superscript_numeric_citations(doc: Document) -> int:
    """Convert [1], [2], [1,2] style citations to superscript."""
    count = 0
    cite_pattern = re.compile(r"\[(\d+(?:[,，\s]\d+)*)\]")

    for para in doc.paragraphs:
        # Check if paragraph contains math elements -- skip those
        if para._element.find(qn("m:oMathPara")) is not None:
            continue
        if para._element.find(qn("m:oMath")) is not None:
            continue

        runs_to_process = []
        for run in para.runs:
            if cite_pattern.search(run.text):
                runs_to_process.append(run)

        for run in runs_to_process:
            text = run.text
            parts = cite_pattern.split(text)
            if not parts:
                continue

            # We need to rebuild this run into multiple runs
            parent = run._element.getparent()
            run_index = list(parent).index(run._element)

            new_elements = []
            last_end = 0
            for m in cite_pattern.finditer(text):
                # Text before citation
                before = text[last_end:m.start()]
                if before:
                    r_before = _make_text_run(before, run)
                    new_elements.append(r_before)

                # Citation text as superscript
                cite_text = m.group(0)
                r_cite = _make_superscript_run(cite_text, run)
                new_elements.append(r_cite)
                count += 1

                last_end = m.end()

            # Remaining text after last citation
            remaining = text[last_end:]
            if remaining:
                r_after = _make_text_run(remaining, run)
                new_elements.append(r_after)

            # Replace original run with new runs
            parent.remove(run._element)
            for i, elem in enumerate(new_elements):
                parent.insert(run_index + i, elem)

    return count
