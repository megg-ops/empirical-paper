"""资产占位符处理 -- 表格渲染、图题检查、占位符替换、图片插入。"""

import logging
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

_FIGURE_PLACEHOLDER = re.compile(r"\[FIGURE:\s*(\w+)\]")
_TABLE_PLACEHOLDER = re.compile(r"\[TABLE:\s*(\w+)\]")


def render_table_asset_as_markdown(tbl_path: str) -> str:
    """将表格资产文件渲染为 Markdown 表格文本。

    支持 csv / xlsx / md 格式。不支持的格式抛出 ValueError。

    Returns:
        Markdown 表格字符串（含表头分隔行）。
    """
    p = Path(tbl_path)
    if not p.exists():
        raise FileNotFoundError(f"表格资产文件不存在: {tbl_path}")

    ext = p.suffix.lower()

    if ext == ".csv":
        import csv
        with open(tbl_path, "r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            rows = [list(row) for row in reader]
        if not rows:
            raise ValueError(f"CSV 表格为空: {tbl_path}")
        # 构建 markdown 表格
        header = rows[0]
        data_rows = rows[1:]
        ncols = len(header)
        # 对齐所有行到相同列数
        for i, row in enumerate(data_rows):
            if len(row) < ncols:
                data_rows[i] = row + [""] * (ncols - len(row))
        lines = ["| " + " | ".join(header) + " |"]
        lines.append("|" + "|".join(["---"] * ncols) + "|")
        for row in data_rows:
            lines.append("| " + " | ".join(row[:ncols]) + " |")
        return "\n".join(lines)

    elif ext in (".xlsx", ".xls"):
        try:
            import openpyxl
        except ImportError:
            raise ImportError("需要 openpyxl 来读取 xlsx 表格，请安装: pip install openpyxl")
        wb = openpyxl.load_workbook(tbl_path, data_only=True, read_only=True)
        ws = wb.active
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(cell) if cell is not None else "" for cell in row])
        wb.close()
        if not rows:
            raise ValueError(f"xlsx 表格为空: {tbl_path}")
        header = rows[0]
        data_rows = rows[1:]
        ncols = len(header)
        lines = ["| " + " | ".join(header) + " |"]
        lines.append("|" + "|".join(["---"] * ncols) + "|")
        for row in data_rows:
            padded = row + [""] * (ncols - len(row)) if len(row) < ncols else row
            lines.append("| " + " | ".join(padded[:ncols]) + " |")
        return "\n".join(lines)

    elif ext == ".md":
        content = p.read_text(encoding="utf-8").strip()
        if not content:
            raise ValueError(f"md 表格文件为空: {tbl_path}")
        # 剥离顶部粗体标题行（如 **变量描述性统计**），标题由 writer 的 caption 提供
        lines = content.split("\n")
        while lines and lines[0].strip().startswith("**") and lines[0].strip().endswith("**"):
            lines.pop(0)
        while lines and not lines[0].strip():
            lines.pop(0)
        return "\n".join(lines)

    else:
        raise ValueError(
            f"不支持的表格资产格式: {ext}（文件: {tbl_path}）。"
            f"仅支持 .csv / .xlsx / .md，不支持的格式不得静默跳过。"
        )


def _check_caption_adjacency(md_text: str) -> list[str]:
    """检查 [TABLE:] 和 [FIGURE:] 占位符周围是否有正确的表题/图题。

    [TABLE: table_XX] 的前一个非空段落必须匹配：^表\\s*\\d+\\s+.+$
    [FIGURE: fig_XX] 的后一个非空段落必须匹配：^图\\s*\\d+\\s+.+$

    Returns:
        错误信息列表，空列表表示全部通过。
    """
    errors = []
    lines = md_text.split("\n")

    for i, line in enumerate(lines):
        stripped = line.strip()

        # Check [TABLE: ...] — previous non-blank line must be a table caption
        tbl_match = re.match(r"^\[TABLE:\s*(\S+?)\s*\]$", stripped)
        if tbl_match:
            tbl_id = tbl_match.group(1)
            prev = ""
            for j in range(i - 1, -1, -1):
                if lines[j].strip():
                    prev = lines[j].strip()
                    break
            if not re.match(r"^表\s*\d+\s+.+$", prev):
                errors.append(
                    f"[TABLE: {tbl_id}] 前一个非空段落必须为表题（格式：'表X 标题内容'），"
                    f"实际为：'{prev or '(无)'}'"
                )

        # Check [FIGURE: ...] — next non-blank line must be a figure caption
        fig_match = re.match(r"^\[FIGURE:\s*(\S+?)\s*\]$", stripped)
        if fig_match:
            fig_id = fig_match.group(1)
            nxt = ""
            for j in range(i + 1, len(lines)):
                if lines[j].strip():
                    nxt = lines[j].strip()
                    break
            if not re.match(r"^图\s*\d+\s+.+$", nxt):
                errors.append(
                    f"[FIGURE: {fig_id}] 后一个非空段落必须为图题（格式：'图X 标题内容'），"
                    f"实际为：'{nxt or '(无)'}'"
                )

    return errors


def replace_asset_placeholders(
    md_text: str,
    assets: dict,
) -> tuple[str, dict]:
    """Replace [FIGURE: id] and [TABLE: id] placeholders in markdown.

    Only replaces content; does NOT add titles or auto-number.
    Title adjacency is validated separately by _check_caption_adjacency().

    Returns (modified_md, stats) where stats has 'figures_replaced', 'tables_replaced',
    'figures_missing', 'tables_missing'.
    """
    stats = {"figures_replaced": 0, "tables_replaced": 0, "figures_missing": [], "tables_missing": []}

    # Build lookup maps
    fig_map = {f["id"]: f for f in assets.get("figures", [])}
    tbl_map = {t["id"]: t for t in assets.get("tables", [])}

    def _replace_fig(m):
        fig_id = m.group(1)
        fig = fig_map.get(fig_id)
        if not fig:
            stats["figures_missing"].append(fig_id)
            return m.group(0)  # leave placeholder as-is
        abs_path = str(Path(fig["path"]).resolve())
        stats["figures_replaced"] += 1
        # alt text 为空，标题由 writer 的 caption（图 X 标题）提供
        return f"![]({abs_path})"

    def _replace_tbl(m):
        tbl_id = m.group(1)
        tbl = tbl_map.get(tbl_id)
        if not tbl:
            stats["tables_missing"].append(tbl_id)
            return m.group(0)  # leave placeholder as-is
        abs_path = str(Path(tbl["path"]).resolve())
        try:
            md_table = render_table_asset_as_markdown(abs_path)
        except (FileNotFoundError, ValueError, ImportError) as exc:
            stats["tables_missing"].append(tbl_id)
            logger.error("TABLE asset render failed for %s: %s", tbl_id, exc)
            return m.group(0)  # leave placeholder as-is; will be caught by validation
        stats["tables_replaced"] += 1
        return md_table

    result = _FIGURE_PLACEHOLDER.sub(_replace_fig, md_text)
    result = _TABLE_PLACEHOLDER.sub(_replace_tbl, result)
    return result, stats


def _validate_assets_embedded(
    md_text: str,
    assets: dict,
    placeholder_stats: dict,
) -> list[str]:
    """Check that all required assets are accounted for.

    Returns list of warning strings (empty = all good).
    """
    warnings = []
    if not assets:
        return warnings

    fig_map = {f["id"]: f for f in assets.get("figures", [])}
    tbl_map = {t["id"]: t for t in assets.get("tables", [])}

    # Check missing figures
    for fig_id in placeholder_stats.get("figures_missing", []):
        fig = fig_map.get(fig_id)
        if fig and fig.get("required", True):
            warnings.append(f"Required figure not found in manifest: {fig_id}")

    # Check missing tables
    for tbl_id in placeholder_stats.get("tables_missing", []):
        tbl = tbl_map.get(tbl_id)
        if tbl and tbl.get("required", True):
            warnings.append(f"Required table not found in manifest: {tbl_id}")

    # Check for leftover placeholders
    leftover_fig = _FIGURE_PLACEHOLDER.findall(md_text)
    if leftover_fig:
        warnings.append(f"Leftover [FIGURE:] placeholders: {leftover_fig}")
    leftover_tbl = _TABLE_PLACEHOLDER.findall(md_text)
    if leftover_tbl:
        warnings.append(f"Leftover [TABLE:] placeholders: {leftover_tbl}")

    # Check for leftover TABLE_ASSET comments (should be consumed by post-processing)
    leftover_assets = re.findall(r'<!--TABLE_ASSET:([^:]+):', md_text)
    if leftover_assets:
        warnings.append(f"Leftover TABLE_ASSET comments: {leftover_assets}")

    # Count placeholders actually present in markdown vs replaced
    replaced_figs = placeholder_stats.get("figures_replaced", 0)
    replaced_tbls = placeholder_stats.get("tables_replaced", 0)
    actual_fig_placeholders = len(_FIGURE_PLACEHOLDER.findall(md_text)) + replaced_figs
    actual_tbl_placeholders = len(_TABLE_PLACEHOLDER.findall(md_text)) + replaced_tbls

    if replaced_figs < actual_fig_placeholders:
        warnings.append(f"Only {replaced_figs}/{actual_fig_placeholders} figure placeholders replaced")
    if replaced_tbls < actual_tbl_placeholders:
        warnings.append(f"Only {replaced_tbls}/{actual_tbl_placeholders} table placeholders replaced")

    return warnings


def insert_images(doc: Document, figures_dir: str) -> int:
    """Insert PNG images at markdown image reference points in the document.
    Looks for paragraphs containing '![alt](filename.png)' patterns.
    """
    if not figures_dir or not os.path.isdir(figures_dir):
        return 0

    figures_path = Path(figures_dir)
    img_pattern = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
    count = 0

    # Collect paragraphs to process (avoid modifying while iterating)
    paras_to_process = []
    for para in doc.paragraphs:
        if img_pattern.search(para.text):
            paras_to_process.append(para)

    for para in paras_to_process:
        full_text = para.text
        matches = list(img_pattern.finditer(full_text))
        if not matches:
            continue

        # For each image reference, try to find the file
        for m in reversed(matches):
            img_filename = m.group(2)
            img_path = figures_path / img_filename

            # Also try with common extensions
            if not img_path.exists():
                for ext in [".png", ".jpg", ".jpeg", ".PNG", ".JPG"]:
                    candidate = figures_path / (img_filename + ext)
                    if candidate.exists():
                        img_path = candidate
                        break

            if not img_path.exists():
                logger.warning("image not found: %s (looked in %s)", img_filename, figures_dir)
                continue

            # Clear the paragraph and insert the image
            # We insert into a new paragraph before this one, then remove text
            # from the original (or clear it if it was only an image ref)
            if full_text.strip() == m.group(0):
                # Entire paragraph is just the image reference
                para.clear()
                run = para.add_run()
                run.add_picture(str(img_path), width=None)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                count += 1
            else:
                # Paragraph has other text; just log a warning
                logger.warning(
                    "image reference '%s' mixed with other text in paragraph, skipping auto-insert",
                    img_filename,
                )

    return count
